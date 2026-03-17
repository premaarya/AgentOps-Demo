from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def set_cell(cell, text, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    return p


def create_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # ---- TITLE PAGE ----
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_heading("Add Agent Button", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_heading("Detailed Functional Documentation", 2)
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")
    meta = doc.add_paragraph("Contract AgentOps - Internal Technical Reference")
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta2 = doc.add_paragraph("Prepared: March 13, 2026")
    meta2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # ---- TABLE OF CONTENTS ----
    doc.add_heading("Table of Contents", 1)
    toc_items = [
        "1. Overview: What Does 'Add Agent' Do?",
        "2. Files Involved (Complete List)",
        "3. Step-by-Step: What Happens When You Click 'Add Agent'",
        "4. The Agent Data Model (What Fields Are Captured)",
        "5. The Modal Form (UI Walkthrough)",
        "6. Validation Logic",
        "7. How the Agent Is Stored In-Memory",
        "8. How the Canvas Re-renders",
        "9. How the Agent Reaches the Backend (Push to Pipeline)",
        "10. Backend Processing: Workflow Route and Registry",
        "11. The Complete Data Flow Diagram",
        "12. Available MCP Tools (Tool Selection Grid)",
        "13. Summary of Key Functions",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # ---- SECTION 1 ----
    doc.add_heading("1. Overview: What Does 'Add Agent' Do?", 1)
    doc.add_paragraph(
        "The 'Add Agent' button allows a user to create a new AI agent and add it to the current "
        "contract processing workflow displayed on the Design Canvas (Tab 1 of the dashboard). "
        "An agent in this system is a specialized AI worker that performs a specific contract-related task, "
        "such as intake classification, clause extraction, compliance checking, or approval routing."
    )
    doc.add_paragraph(
        "When you click 'Add Agent', a modal dialog opens with a form where you define the agent's name, "
        "role, AI model, MCP tools, execution constraints, and visual properties. After filling in the form "
        "and clicking 'Add Agent' in the modal, the new agent is added to the in-memory workflow state, "
        "the canvas re-renders to show the agent as a new card in the pipeline, and the inventory bar updates "
        "with new totals."
    )
    doc.add_paragraph(
        "The agent is NOT automatically saved to the backend. It lives in the browser's in-memory state "
        "until the user explicitly clicks 'Save' (persists to localStorage) or 'Push to Pipeline' "
        "(saves to localStorage AND sends to the backend gateway via REST API)."
    )

    # ---- SECTION 2 ----
    doc.add_heading("2. Files Involved (Complete List)", 1)
    doc.add_paragraph(
        "The Add Agent functionality spans across the frontend UI and the backend gateway. "
        "Here is every file involved, from the button click to backend persistence:"
    )

    doc.add_heading("Frontend Files (Browser)", 2)
    files_fe = [
        ("ui/index.html", "The main HTML page. Contains the Design Canvas section with placeholder containers: "
         "#designer-toolbar (where the Add Agent button is rendered), #designer-canvas (where agent cards appear), "
         "#designer-inventory (summary bar), and #designer-modal (overlay container for the agent form). "
         "These containers are empty in the HTML -- they are populated dynamically by JavaScript."),
        ("ui/workflow-designer.js", "THE CORE FILE. Contains the entire WorkflowDesigner module as an IIFE "
         "(Immediately Invoked Function Expression). This single file handles: the Add Agent button rendering, "
         "the modal form, agent data validation, in-memory state management, canvas rendering, drag-and-drop, "
         "save/load, export, and push-to-pipeline. The addAgent() function (line 724) and saveAgent() function "
         "(line 893) are the two key functions for this feature."),
        ("ui/styles.css", "CSS styles for the modal (.designer-modal, .designer-modal-content, .designer-form-row, "
         ".designer-form-group, .designer-tool-grid), the agent cards (.agent-card, .designer-agent-card), "
         "the toolbar (.designer-toolbar), and the canvas (.designer-canvas-wrapper, .agent-pipeline)."),
        ("ui/app.js", "The main dashboard application script. It handles tab navigation and downstream workflow "
         "integration. When a workflow is pushed to the pipeline, app.js receives a 'workflow-activated' event "
         "and propagates the workflow to the Test, Deploy, and Live tabs. Not directly involved in Add Agent, "
         "but consumes the output."),
        ("ui/api.js", "The API integration layer. Provides the apiCall() helper and GATEWAY_URL constant used "
         "when Push to Pipeline sends the workflow (including the new agent) to the backend."),
    ]
    t_fe = doc.add_table(rows=len(files_fe) + 1, cols=2)
    t_fe.style = "Light Grid Accent 1"
    set_cell(t_fe.rows[0].cells[0], "File", bold=True)
    set_cell(t_fe.rows[0].cells[1], "Purpose", bold=True)
    for i, (f, desc) in enumerate(files_fe, 1):
        set_cell(t_fe.rows[i].cells[0], f)
        set_cell(t_fe.rows[i].cells[1], desc)

    doc.add_paragraph("")
    doc.add_heading("Backend Files (Node.js Gateway)", 2)
    files_be = [
        ("gateway/src/routes/workflows.ts", "REST API route handler. Defines POST /api/v1/workflows (saves a workflow "
         "definition including all agents) and POST /api/v1/workflows/:id/activate (sets it as active). "
         "These endpoints are called when the user clicks 'Push to Pipeline' after adding an agent."),
        ("gateway/src/services/workflowRegistry.ts", "The core backend service. Manages workflow persistence in "
         "a JSON file (data/workflows/definitions.json), validates workflow input, generates workflow packages "
         "with runtime agent bindings, maps agents to contract lifecycle stages, and resolves role-specific "
         "assets (prompts, configs, schemas)."),
        ("gateway/src/config.ts", "Application configuration. Provides appConfig with the dataDir path where "
         "workflow definitions are stored on disk."),
        ("gateway/src/stores/jsonStore.ts", "Generic JSON file persistence helper. The workflowRegistry uses "
         "JsonStore<WorkflowDefinition> to read/write the definitions.json file atomically."),
        ("config/stages/contract-lifecycle.json", "The contract stage catalog. When a workflow is activated, "
         "the registry maps each agent to contract lifecycle stages (Intake, Drafting, Review, etc.) using "
         "this catalog as the reference."),
        ("config/agents/*.yaml", "Per-role agent configuration files (intake-agent.yaml, compliance-agent.yaml, "
         "etc.). These are referenced in the workflow package as declarative agent configs."),
        ("prompts/*.md", "System prompts for each agent role (intake-system.md, compliance-system.md, etc.). "
         "Referenced in the workflow package when an agent is mapped to a known role."),
        ("data/workflows/definitions.json", "The on-disk JSON file where all saved workflow definitions "
         "(including their agents) are persisted. Created automatically by the JsonStore."),
    ]
    t_be = doc.add_table(rows=len(files_be) + 1, cols=2)
    t_be.style = "Light Grid Accent 1"
    set_cell(t_be.rows[0].cells[0], "File", bold=True)
    set_cell(t_be.rows[0].cells[1], "Purpose", bold=True)
    for i, (f, desc) in enumerate(files_be, 1):
        set_cell(t_be.rows[i].cells[0], f)
        set_cell(t_be.rows[i].cells[1], desc)
    doc.add_page_break()

    # ---- SECTION 3 ----
    doc.add_heading("3. Step-by-Step: What Happens When You Click 'Add Agent'", 1)
    doc.add_paragraph(
        "This section provides an exact walkthrough of what happens from the moment the user clicks "
        "the button, through each function call, to the final rendering of the new agent card."
    )

    steps = [
        ("Step 1: User Clicks the '+ Add Agent' Button",
         "The button is located in the Design Canvas toolbar (#designer-toolbar). It is rendered dynamically "
         "by the renderToolbar() function in workflow-designer.js. The button HTML is:\n\n"
         '<button class="btn btn-outline" onclick="WorkflowDesigner.addAgent()" title="Add a new agent">\n'
         "  + Add Agent\n</button>\n\n"
         "When clicked, it calls WorkflowDesigner.addAgent()."),
        ("Step 2: addAgent() Creates a Blank Agent Object",
         "The addAgent() function (line 724 of workflow-designer.js) creates a new agent data object with "
         "default values:\n\n"
         "- id: Auto-generated as 'agent-N' where N is an incrementing counter (nextAgentId++)\n"
         "- name: Empty string (user must fill in)\n"
         "- role: Empty string (user must fill in)\n"
         "- icon: Empty string (user must fill in, 1-2 characters)\n"
         "- model: 'GPT-5.4' (default selection)\n"
         "- tools: Empty array (user selects from checkbox grid)\n"
         "- boundary: Empty string (optional constraint)\n"
         "- output: Empty string (optional description)\n"
         "- kind: 'agent' (default execution role)\n"
         "- color: Next color from AGENT_COLORS palette (cycles through 12 colors)\n"
         "- stage: Set to currentWorkflow.agents.length (appended as last stage)\n"
         "- lane: 0 (first lane)\n"
         "- order: Set to currentWorkflow.agents.length"),
        ("Step 3: openAgentModal() Renders the Form",
         "addAgent() calls openAgentModal(newAgent, true) where 'true' indicates this is a new agent "
         "(not editing an existing one). This function:\n\n"
         "1. Finds the #designer-modal container in the DOM\n"
         "2. Builds the complete HTML form with all fields\n"
         "3. Populates the MCP tool selection grid from AVAILABLE_TOOLS constant\n"
         "4. Sets the modal title to 'Add New Agent' (vs 'Edit Agent' for existing)\n"
         "5. Sets the submit button text to 'Add Agent' (vs 'Save Changes')\n"
         "6. Adds the 'visible' CSS class to show the modal with a backdrop overlay"),
        ("Step 4: User Fills In the Form",
         "The user enters data in the modal form fields:\n"
         "- Agent Name (text input, required)\n"
         "- Icon (text input, 1-2 characters, required)\n"
         "- Model (dropdown: GPT-5.4, GPT-4o-mini, GPT-4.1, etc.)\n"
         "- Role/Responsibility (textarea, required)\n"
         "- Tools (checkbox grid organized by MCP server)\n"
         "- Boundary Constraint (text input, optional)\n"
         "- Output Description (text input, optional)\n"
         "- Execution Role (dropdown: Agent, Orchestrator, Human Checkpoint, Merge Node)\n"
         "- Stage number (numeric input)\n"
         "- Lane number (numeric input)\n"
         "- Agent Color (visual color picker with 12 swatches)"),
        ("Step 5: User Clicks 'Add Agent' in the Modal",
         "The modal footer has the submit button:\n"
         '<button class="btn btn-primary" onclick="WorkflowDesigner.saveAgent(agentId, true)">\n'
         "  Add Agent\n</button>\n\n"
         "This calls saveAgent(agentId, isNew=true)."),
        ("Step 6: saveAgent() Reads, Validates, and Stores",
         "The saveAgent() function (line 893) performs these steps:\n\n"
         "1. READS all form values from DOM elements using document.getElementById()\n"
         "2. GATHERS selected tools by querying all checked checkboxes in #modal-tool-grid\n"
         "3. VALIDATES required fields: name, icon, and role must not be empty. If any are missing, "
         "an alert() dialog is shown and the function returns without saving.\n"
         "4. CONSTRUCTS the agentData object from all collected values\n"
         "5. Since isNew=true: sets agentData.order and PUSHES it into currentWorkflow.agents array\n"
         "6. NORMALIZES the workflow via normalizeWorkflow() which re-sorts agents by stage/lane/order, "
         "compacts stage numbers, and re-indexes lane assignments\n"
         "7. CLOSES the modal by removing the 'visible' class\n"
         "8. Calls render() to refresh the entire canvas"),
        ("Step 7: render() Refreshes All Three Areas",
         "The render() function calls three sub-renderers:\n\n"
         "1. renderToolbar() -- Re-renders the top bar (agent count in buttons stays current)\n"
         "2. renderCanvas() -- Rebuilds the entire pipeline visualization with all agent cards, "
         "including the new one. Each agent becomes a draggable card with edit/delete/drag buttons.\n"
         "3. renderInventory() -- Updates the bottom summary bar with new totals: "
         "Total Agents, Stages, MCP Tools, Model, Pipeline type, and Design Validation status."),
        ("Step 8: Toast Notification",
         'A small toast message appears at the bottom of the screen: "[Agent Name] added to workflow". '
         "It automatically disappears after 2.5 seconds."),
    ]
    for title, desc in steps:
        p_step = doc.add_paragraph()
        p_step.add_run(title + "\n").bold = True
        p_step.add_run(desc)
    doc.add_page_break()

    # ---- SECTION 4 ----
    doc.add_heading("4. The Agent Data Model (What Fields Are Captured)", 1)
    doc.add_paragraph(
        "Every agent in the system is represented by a JavaScript object with the following properties. "
        "This is the complete schema:"
    )

    fields = [
        ("id", "string", "Auto-generated unique identifier (e.g., 'agent-7'). Used internally for lookups, drag-drop, edit, and delete operations."),
        ("name", "string", "Human-readable agent name (e.g., 'Intake Agent'). Required. Displayed on the card and in all references."),
        ("role", "string", "Description of what the agent does (e.g., 'Classify contracts by type'). Required. Appears in the card's Role section."),
        ("icon", "string", "1-2 character icon shown in the colored circle on the card (e.g., 'I', 'C', 'DR'). Required."),
        ("model", "string", "AI model selection. Options: GPT-5.4, GPT-4o-mini, GPT-4.1, GPT-4.1-mini, GPT-4.1-nano, o3-mini, o4-mini. Default: GPT-5.4."),
        ("tools", "string[]", "Array of MCP tool names the agent is allowed to call. Selected from AVAILABLE_TOOLS organized by server. Can be empty."),
        ("boundary", "string", "Execution constraint text (e.g., 'Classify only', 'Flag only'). Optional. Tells the agent what it should NOT do."),
        ("output", "string", "Expected output description (e.g., 'Contract classification and metadata'). Optional."),
        ("kind", "string", "Execution role. One of: 'agent' (standard), 'orchestrator' (fan-out controller), 'human' (HITL checkpoint), 'merge' (fan-in collector). Default: 'agent'."),
        ("color", "string", "Hex color for the card border and icon background. Selected from 12-color palette. Auto-assigned based on position."),
        ("stage", "number", "Pipeline stage number (0-based). Determines vertical grouping on the canvas. Auto-assigned as the next stage."),
        ("lane", "number", "Lane within a stage (0-based). Multiple agents in the same stage with different lanes create a parallel layout."),
        ("order", "number", "Global ordering index. Used for sorting and normalization."),
    ]
    t_fields = doc.add_table(rows=len(fields) + 1, cols=3)
    t_fields.style = "Light Grid Accent 1"
    set_cell(t_fields.rows[0].cells[0], "Field", bold=True)
    set_cell(t_fields.rows[0].cells[1], "Type", bold=True)
    set_cell(t_fields.rows[0].cells[2], "Description", bold=True)
    for i, (field, ftype, desc) in enumerate(fields, 1):
        set_cell(t_fields.rows[i].cells[0], field)
        set_cell(t_fields.rows[i].cells[1], ftype)
        set_cell(t_fields.rows[i].cells[2], desc)
    doc.add_page_break()

    # ---- SECTION 5 ----
    doc.add_heading("5. The Modal Form (UI Walkthrough)", 1)
    doc.add_paragraph(
        "When openAgentModal() is called, it dynamically generates a full-screen modal overlay. "
        "The modal has three sections:"
    )
    doc.add_heading("Header", 2)
    doc.add_paragraph(
        "Shows the title ('Add New Agent' for new agents, 'Edit Agent' for existing ones) and a close button (X)."
    )
    doc.add_heading("Body (The Form)", 2)
    doc.add_paragraph("The form is organized into rows of fields:")

    form_rows = [
        ("Row 1: Identity", "Three fields side by side:\n"
         "- Agent Name (flex:2, text input, placeholder: 'e.g. Intake Agent')\n"
         "- Icon (flex:0 0 80px, text input, max 2 chars, centered, bold, 18px font)\n"
         "- Model (flex:1, dropdown select with 7 model options)"),
        ("Row 2: Role", "Full-width textarea (2 rows) where the user describes the agent's responsibility. "
         "Placeholder: 'Describe what this agent does...'"),
        ("Row 3: Tools Grid", "A visual grid organized by MCP server. Each server (8 total) shows as a labeled "
         "group with checkboxes for its tools. The user checks the tools this agent should have access to. "
         "Example: under 'contract-intake-mcp', options are upload_contract, classify_document, extract_metadata."),
        ("Row 4: Constraints", "Two fields side by side:\n"
         "- Boundary Constraint (text input, placeholder: 'e.g. Classify only')\n"
         "- Output Description (text input, placeholder: 'e.g. Classification and metadata')"),
        ("Row 5: Execution", "Three fields side by side:\n"
         "- Execution Role (dropdown: Agent, Orchestrator, Human Checkpoint, Merge Node)\n"
         "- Stage (number input, min 0)\n"
         "- Lane (number input, min 0)"),
        ("Row 6: Color", "A row of 12 colored squares. Click one to select the agent's card color. "
         "The selected swatch gets a 'selected' border. Colors cycle through: blue, teal, purple, orange, "
         "red, light blue, green, violet, cyan, amber, pink, and dark cyan."),
    ]
    for title, desc in form_rows:
        p = doc.add_paragraph()
        p.add_run(title + "\n").bold = True
        p.add_run(desc)

    doc.add_heading("Footer", 2)
    doc.add_paragraph(
        "Two buttons:\n- Cancel: Closes the modal without saving (also triggered by clicking the backdrop or pressing Escape)\n"
        "- Add Agent (primary blue button): Calls saveAgent() to validate and commit the new agent"
    )

    # ---- SECTION 6 ----
    doc.add_heading("6. Validation Logic", 1)
    doc.add_paragraph("Validation happens at two levels:")

    doc.add_heading("6.1 Immediate Validation (On Save)", 2)
    doc.add_paragraph(
        "When the user clicks 'Add Agent' in the modal, saveAgent() checks three required fields:\n\n"
        "1. name -- Must not be empty. Error: 'Agent name is required.'\n"
        "2. icon -- Must not be empty. Error: 'Agent icon is required (1-2 characters).'\n"
        "3. role -- Must not be empty. Error: 'Agent role/responsibility is required.'\n\n"
        "If any of these are empty, an alert() dialog is shown and the function returns without saving. "
        "The modal stays open so the user can fix the issue."
    )

    doc.add_heading("6.2 Workflow-Level Validation (On Render)", 2)
    doc.add_paragraph(
        "After the agent is added and render() is called, the renderInventory() function runs "
        "validateWorkflow() which performs comprehensive design validation. This produces findings "
        "displayed in the Design Validation section of the inventory bar. Checks include:\n\n"
        "- Each agent must have a name, role, and model\n"
        "- Agents should have a boundary and output defined (warnings if missing)\n"
        "- Agents should have at least one tool assigned (warning if none)\n"
        "- Tools must exist in the known MCP tool registry\n"
        "- Workflow type constraints are enforced:\n"
        "  -- Sequential/Conditional: No parallel agents in any stage\n"
        "  -- Sequential HITL: Must have at least one human checkpoint\n"
        "  -- Parallel: Must start with a single orchestrator\n"
        "  -- Fan-out: Must start with orchestrator and end with merge node\n\n"
        "Findings are classified as errors (block save/deploy), warnings (allow but flag), or info. "
        "The validation status badge shows: 'Ready' (green), 'Needs Review' (yellow), or 'Blocked' (red)."
    )

    # ---- SECTION 7 ----
    doc.add_heading("7. How the Agent Is Stored In-Memory", 1)
    doc.add_paragraph(
        "The WorkflowDesigner module uses an IIFE (Immediately Invoked Function Expression) closure pattern. "
        "The workflow state is stored in a private variable called 'currentWorkflow' inside the closure. "
        "This variable is a plain JavaScript object with this structure:"
    )
    add_code_block(doc,
        'let currentWorkflow = {\n'
        '  id: "wf-m2abc123",\n'
        '  name: "Contract Lifecycle Pipeline",\n'
        '  type: "sequential-hitl",\n'
        '  agents: [\n'
        '    { id: "agent-1", name: "Intake Agent", role: "...", ... },\n'
        '    { id: "agent-2", name: "Drafting Agent", role: "...", ... },\n'
        '    // ... more agents ...\n'
        '    { id: "agent-7", name: "NEW AGENT", role: "...", ... }  // <-- just added\n'
        '  ],\n'
        '  createdAt: "2026-03-13T...",\n'
        '  updatedAt: "2026-03-13T..."\n'
        '};'
    )
    doc.add_paragraph(
        "When saveAgent() runs with isNew=true, it pushes the new agent object onto the agents array. "
        "Then normalizeWorkflow() is called, which:\n\n"
        "1. Re-sorts all agents by (stage, lane, order, name)\n"
        "2. Compacts stage numbers (removes gaps)\n"
        "3. Re-indexes lanes within each stage\n"
        "4. Re-indexes global order numbers\n\n"
        "This normalization ensures the workflow is always in a consistent state, even if the user "
        "entered a stage number that creates a gap."
    )
    doc.add_paragraph(
        "IMPORTANT: At this point, the agent exists ONLY in browser memory. It is NOT persisted "
        "anywhere until the user takes one of these actions:\n\n"
        "- 'Save' button: Persists to browser localStorage under key 'agentops-workflows'\n"
        "- 'Push to Pipeline' button: Persists to localStorage AND sends to backend via REST API"
    )

    # ---- SECTION 8 ----
    doc.add_heading("8. How the Canvas Re-renders", 1)
    doc.add_paragraph(
        "After the agent is added, render() calls renderCanvas(), which completely rebuilds the "
        "pipeline visualization. Here is the rendering process:"
    )
    render_steps = [
        ("1. Normalize", "normalizeWorkflow() ensures agents are sorted and stages are compact."),
        ("2. Get Stages", "getWorkflowStages() groups agents into stage objects. Each stage has a stage number, "
         "an array of agents, and an isParallel flag (true if more than one agent in the stage)."),
        ("3. Build Stage Flow", "renderStageLayout() creates a horizontal flow of stage containers connected by arrow dividers."),
        ("4. Render Each Stage", "renderStage() creates a stage container with a label ('Stage 1', 'Stage 2 - Parallel', etc.) "
         "and either a sequential or parallel grid layout inside."),
        ("5. Render Agent Cards", "renderAgentCard() creates each agent's visual card with: edit/delete/drag buttons, "
         "colored icon circle, name, kind badge (Agent/Orchestrator/Human/Merge), stage and lane badges, "
         "model name, role text, tools list, boundary text, and output text."),
        ("6. Bind Drag Events", "bindDragEvents() attaches HTML5 drag-and-drop handlers to all agent cards, "
         "allowing the user to reorder agents by dragging."),
    ]
    for title, desc in render_steps:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title + ": ").bold = True
        p.add_run(desc)
    doc.add_page_break()

    # ---- SECTION 9 ----
    doc.add_heading("9. How the Agent Reaches the Backend (Push to Pipeline)", 1)
    doc.add_paragraph(
        "After adding an agent, the user must click 'Push to Pipeline' to send the workflow (with all "
        "agents, including the new one) to the backend gateway. Here is the exact sequence:"
    )
    push_steps = [
        ("1. pushToPipeline() is called", "This function first calls saveWorkflow(), which validates "
         "the workflow, assigns an ID if needed, saves to localStorage, and returns true on success. "
         "If validation fails (any errors), the push is aborted."),
        ("2. POST /api/v1/workflows", "The function sends a POST request to the gateway with the "
         "full workflow definition:\n\n"
         "  URL: http://localhost:8000/api/v1/workflows\n"
         "  Method: POST\n"
         "  Body: { id, name, type, agents: [...] }\n\n"
         "The gateway's workflows route handler receives this, validates it, and calls "
         "saveWorkflowDefinition() from the workflowRegistry service, which persists it to "
         "data/workflows/definitions.json."),
        ("3. POST /api/v1/workflows/:id/activate", "After the save succeeds, a second POST request "
         "activates this workflow:\n\n"
         "  URL: http://localhost:8000/api/v1/workflows/{workflowId}/activate\n"
         "  Method: POST\n\n"
         "The gateway calls activateWorkflowDefinition() which:\n"
         "  a) Marks this workflow as active in the definitions store\n"
         "  b) Generates a full WorkflowPackage with:\n"
         "     - RuntimeAgentBindings for each agent (with role keys, declarative configs)\n"
         "     - ContractStageMap mapping agents to lifecycle stages\n"
         "     - HITL policy, model policy, mode policy\n"
         "  c) Saves the package to data/runtime/active-workflow.json\n"
         "  d) Returns the activation result to the frontend"),
        ("4. 'workflow-activated' Event", "The frontend dispatches a CustomEvent('workflow-activated') "
         "that app.js listens for. This propagates the new active workflow to the Test, Deploy, and Live tabs, "
         "updating their content to reflect the new workflow including the added agent."),
    ]
    for title, desc in push_steps:
        p = doc.add_paragraph()
        p.add_run(title + "\n").bold = True
        p.add_run(desc)

    # ---- SECTION 10 ----
    doc.add_heading("10. Backend Processing: Workflow Route and Registry", 1)
    doc.add_paragraph(
        "When the workflow arrives at the backend, the following processing happens:"
    )
    doc.add_heading("10.1 Validation (workflows.ts)", 2)
    doc.add_paragraph(
        "The POST /api/v1/workflows route validates the incoming request body:\n"
        "- name, type, and agents must be present\n"
        "- If validation fails, returns HTTP 400 with error details\n"
        "- If an existing workflow with the same ID exists, it is updated; otherwise, a new one is created"
    )
    doc.add_heading("10.2 Persistence (workflowRegistry.ts)", 2)
    doc.add_paragraph(
        "saveWorkflowDefinition() in the workflowRegistry service:\n"
        "- Assigns a UUID if no ID is provided\n"
        "- Sets createdAt and updatedAt timestamps\n"
        "- Writes the complete workflow definition (including all agent objects) to a JsonStore backed by "
        "data/workflows/definitions.json\n"
        "- The JsonStore uses atomic file writes (write to temp, then rename) for safety"
    )
    doc.add_heading("10.3 Activation (workflowRegistry.ts)", 2)
    doc.add_paragraph(
        "activateWorkflowDefinition() transforms the design-time workflow into a runtime-ready package:\n\n"
        "1. Each agent gets a runtime_role_key derived from its properties (name, role, tools, boundary)\n"
        "2. Based on the role_key, declarative assets are resolved (agent config YAML, system prompt MD, "
        "output schema JSON) -- for example, an agent with role_key 'intake' gets linked to:\n"
        "   - config/agents/intake-agent.yaml\n"
        "   - prompts/intake-system.md\n"
        "   - config/schemas/intake-result.json\n"
        "3. The ContractStageMap is generated by matching agents to the contract lifecycle catalog "
        "(config/stages/contract-lifecycle.json), which defines stages like Intake, Drafting, Review, "
        "Compliance, Negotiation, Approval, Execution, Obligations, Renewal, Analytics\n"
        "4. The complete WorkflowPackage is written to data/runtime/active-workflow.json\n"
        "5. A historical copy is saved in data/runtime/packages/"
    )
    doc.add_paragraph(
        "The WorkflowAgent TypeScript interface in the registry defines these fields:\n"
        "id, name, role, icon, model, tools, boundary, output, color, kind, stage, lane, order\n\n"
        "This matches exactly what the frontend sends -- there is a 1:1 correspondence between "
        "the frontend agent object and the backend WorkflowAgent interface."
    )
    doc.add_page_break()

    # ---- SECTION 11 ----
    doc.add_heading("11. The Complete Data Flow Diagram", 1)
    doc.add_paragraph("Here is the complete data flow from button click to backend persistence:")
    doc.add_paragraph("")
    add_code_block(doc,
        "User clicks '+ Add Agent' (toolbar)\n"
        "       |\n"
        "       v\n"
        "WorkflowDesigner.addAgent()         [workflow-designer.js line 724]\n"
        "  Creates blank agent object\n"
        "       |\n"
        "       v\n"
        "openAgentModal(agent, isNew=true)    [workflow-designer.js line 757]\n"
        "  Renders form in #designer-modal\n"
        "       |\n"
        "       v\n"
        "User fills form and clicks 'Add Agent'\n"
        "       |\n"
        "       v\n"
        "WorkflowDesigner.saveAgent(id, true) [workflow-designer.js line 893]\n"
        "  1. Read form values from DOM\n"
        "  2. Validate name/icon/role\n"
        "  3. Push to currentWorkflow.agents[]\n"
        "  4. normalizeWorkflow()\n"
        "  5. closeModal()\n"
        "  6. render()\n"
        "       |\n"
        "       v\n"
        "render()                              [workflow-designer.js]\n"
        "  renderToolbar() + renderCanvas() + renderInventory()\n"
        "  Agent card now visible on canvas\n"
        "       |\n"
        "       | (User clicks 'Push to Pipeline')\n"
        "       v\n"
        "pushToPipeline()                     [workflow-designer.js line 1091]\n"
        "  1. saveWorkflow() -> localStorage\n"
        "  2. POST /api/v1/workflows -> gateway\n"
        "       |\n"
        "       v\n"
        "workflows.ts route handler            [gateway/src/routes/workflows.ts]\n"
        "  -> saveWorkflowDefinition()\n"
        "       |\n"
        "       v\n"
        "workflowRegistry.ts                   [gateway/src/services/workflowRegistry.ts]\n"
        "  -> JsonStore.save() -> data/workflows/definitions.json\n"
        "       |\n"
        "       v\n"
        "POST /api/v1/workflows/:id/activate\n"
        "  -> activateWorkflowDefinition()\n"
        "  -> Generate WorkflowPackage\n"
        "  -> Map agents to contract stages\n"
        "  -> Write data/runtime/active-workflow.json\n"
        "       |\n"
        "       v\n"
        "Frontend receives activation response\n"
        "  -> Dispatches 'workflow-activated' event\n"
        "  -> Test/Deploy/Live tabs update"
    )
    doc.add_page_break()

    # ---- SECTION 12 ----
    doc.add_heading("12. Available MCP Tools (Tool Selection Grid)", 1)
    doc.add_paragraph(
        "When adding an agent, the user can select which MCP tools the agent is allowed to call. "
        "These tools are organized by MCP server. Here is the complete list:"
    )

    tools = [
        ("contract-intake-mcp", "upload_contract, classify_document, extract_metadata"),
        ("contract-extraction-mcp", "extract_clauses, identify_parties, extract_dates_values"),
        ("contract-compliance-mcp", "check_policy, flag_risk, get_policy_rules"),
        ("contract-workflow-mcp", "route_approval, escalate_to_human, notify_stakeholder"),
        ("contract-audit-mcp", "get_audit_log, create_audit_entry"),
        ("contract-eval-mcp", "run_evaluation, get_baseline"),
        ("contract-drift-mcp", "detect_drift, model_swap_analysis"),
        ("contract-feedback-mcp", "submit_feedback, optimize_feedback"),
    ]
    t_tools = doc.add_table(rows=len(tools) + 1, cols=2)
    t_tools.style = "Light Grid Accent 1"
    set_cell(t_tools.rows[0].cells[0], "MCP Server", bold=True)
    set_cell(t_tools.rows[0].cells[1], "Available Tools", bold=True)
    for i, (server, tool_list) in enumerate(tools, 1):
        set_cell(t_tools.rows[i].cells[0], server)
        set_cell(t_tools.rows[i].cells[1], tool_list)

    doc.add_paragraph("")
    doc.add_paragraph(
        "The tool grid renders as checkboxes grouped by server. Each checkbox stores the tool name "
        "and a data-server attribute. When saveAgent() runs, it queries all checked checkboxes and "
        "collects their values into the agent's tools array."
    )

    # ---- SECTION 13 ----
    doc.add_heading("13. Summary of Key Functions", 1)
    doc.add_paragraph(
        "Here is a quick reference of every function involved in the Add Agent feature, "
        "the file it lives in, and what it does:"
    )

    funcs = [
        ("WorkflowDesigner.addAgent()", "workflow-designer.js", "Creates blank agent object, opens the modal form"),
        ("openAgentModal(agent, isNew)", "workflow-designer.js", "Builds and displays the full agent editor modal"),
        ("WorkflowDesigner.saveAgent(id, isNew)", "workflow-designer.js", "Reads form, validates, pushes to workflow, re-renders"),
        ("normalizeWorkflow(workflow)", "workflow-designer.js", "Sorts agents, compacts stages, re-indexes lanes and order"),
        ("render()", "workflow-designer.js", "Calls renderToolbar + renderCanvas + renderInventory"),
        ("renderToolbar()", "workflow-designer.js", "Renders the top bar with buttons including '+ Add Agent'"),
        ("renderCanvas()", "workflow-designer.js", "Rebuilds the full pipeline visualization with agent cards"),
        ("renderAgentCard(agent)", "workflow-designer.js", "Generates HTML for a single agent card"),
        ("renderInventory()", "workflow-designer.js", "Updates the summary bar with totals and validation results"),
        ("validateWorkflow()", "workflow-designer.js", "Runs comprehensive design checks, returns findings"),
        ("closeModal()", "workflow-designer.js", "Hides the modal and clears its content"),
        ("showToast(message)", "workflow-designer.js", "Shows a temporary notification at the bottom"),
        ("selectColor(el, color)", "workflow-designer.js", "Handles color swatch selection in the modal"),
        ("escapeHtml(str)", "workflow-designer.js", "Sanitizes strings before inserting into HTML (XSS prevention)"),
        ("pushToPipeline()", "workflow-designer.js", "Saves + POSTs workflow to gateway + activates it"),
        ("saveWorkflow()", "workflow-designer.js", "Validates and persists to localStorage"),
        ("apiCall(method, path, body)", "api.js", "Generic fetch wrapper for gateway REST calls"),
        ("POST /api/v1/workflows", "workflows.ts", "Backend route: validates and saves workflow definition"),
        ("POST /api/v1/workflows/:id/activate", "workflows.ts", "Backend route: activates workflow, generates package"),
        ("saveWorkflowDefinition()", "workflowRegistry.ts", "Persists workflow to JSON file on disk"),
        ("activateWorkflowDefinition()", "workflowRegistry.ts", "Generates runtime package, maps agents to stages"),
    ]
    t_funcs = doc.add_table(rows=len(funcs) + 1, cols=3)
    t_funcs.style = "Light Grid Accent 1"
    set_cell(t_funcs.rows[0].cells[0], "Function / Endpoint", bold=True)
    set_cell(t_funcs.rows[0].cells[1], "File", bold=True)
    set_cell(t_funcs.rows[0].cells[2], "Purpose", bold=True)
    for i, (func, file, purpose) in enumerate(funcs, 1):
        set_cell(t_funcs.rows[i].cells[0], func)
        set_cell(t_funcs.rows[i].cells[1], file)
        set_cell(t_funcs.rows[i].cells[2], purpose)

    # Save
    doc_path = "C:\\Users\\pfernandes\\AgentOps\\AgentOps\\AddAgent.docx"
    doc.save(doc_path)
    print(f"Document created successfully at {doc_path}")


if __name__ == "__main__":
    create_doc()
