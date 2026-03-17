"""
Generate a Word document explaining the backend behaviour of the
Test, Deploy and Live tabs in Real Mode for AgentOps.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2B579A")

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "E8EDF4")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def build_document():
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    # ========================================================
    # TITLE PAGE
    # ========================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AgentOps - Contract Management Platform")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Backend Deep-Dive: Test, Deploy & Live Tabs in Real Mode")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("March 2026  |  Internal Technical Documentation").font.size = Pt(11)

    doc.add_page_break()

    # ========================================================
    # TABLE OF CONTENTS placeholder
    # ========================================================
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "1. Introduction & Architecture Overview\n"
        "2. Mode Toggle: Real vs Simulated\n"
        "3. Test Tab - Backend Walkthrough\n"
        "4. Deploy Tab - Backend Walkthrough\n"
        "5. Live Tab - Backend Walkthrough\n"
        "6. API Reference Summary\n"
        "7. Data Flow Diagrams"
    )
    doc.add_page_break()

    # ========================================================
    # 1. INTRODUCTION
    # ========================================================
    doc.add_heading("1. Introduction & Architecture Overview", level=1)
    doc.add_paragraph(
        "AgentOps is a multi-agent contract management platform built on a microservices "
        "architecture. The system consists of three main layers:"
    )
    doc.add_paragraph(
        "UI Layer - A single-page application (HTML/CSS/JS) that provides three operational "
        "tabs: Test, Deploy, and Live. Each tab calls a different set of backend APIs depending "
        "on whether the dashboard is in Real Mode or Simulated Mode.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Gateway Layer - A Fastify-based Node.js HTTP server that exposes RESTful endpoints "
        "and a WebSocket stream. The gateway orchestrates all backend logic: workflow validation, "
        "Foundry deployment, and contract processing pipeline execution.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "MCP Server Layer - Eight specialised Model Context Protocol (MCP) servers, each exposing "
        "domain-specific tools (intake, extraction, compliance, workflow, audit, evaluation, "
        "drift detection, and feedback).",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Agent Layer - Six AI agents (Intake, Drafting/Extraction, Review, Compliance, "
        "Negotiation, Approval) that call Azure OpenAI via an LLM adapter. Each agent loads a "
        "system prompt from the prompts/ directory and returns structured JSON.",
        style="List Bullet"
    )

    doc.add_heading("System Startup (start.ts)", level=2)
    doc.add_paragraph(
        "On launch, start.ts spawns all eight MCP servers in parallel (ports 9001-9008), "
        "polls each /health endpoint until ready, then starts the Gateway server. The Gateway "
        "serves both the REST API and the static UI files. A consolidated health check confirms "
        "all services are online before the system is considered ready."
    )

    add_styled_table(doc,
        ["MCP Server", "Port", "Purpose"],
        [
            ["contract-intake-mcp", "9001", "Contract upload, classification, metadata extraction tools"],
            ["contract-extraction-mcp", "9002", "Clause extraction, party identification, date/value extraction"],
            ["contract-compliance-mcp", "9003", "Policy checking, risk flagging, policy rule retrieval"],
            ["contract-workflow-mcp", "9004", "Workflow orchestration, stage management"],
            ["contract-audit-mcp", "9005", "Audit logging, audit trail retrieval"],
            ["contract-eval-mcp", "9006", "Agent evaluation, quality scoring"],
            ["contract-drift-mcp", "9007", "Drift detection, monitoring"],
            ["contract-feedback-mcp", "9008", "Human feedback collection and routing"],
        ],
        col_widths=[5, 2, 10],
    )

    doc.add_page_break()

    # ========================================================
    # 2. MODE TOGGLE
    # ========================================================
    doc.add_heading("2. Mode Toggle: Real vs Simulated", level=1)
    doc.add_paragraph(
        "The dashboard exposes a mode toggle switch that sends a POST request to the gateway:"
    )
    doc.add_paragraph(
        "Endpoint: POST /api/v1/mode\n"
        'Request body: { "mode": "live" } or { "mode": "simulated" }',
        style="List Bullet"
    )
    doc.add_paragraph(
        "When set to \"live\" (Real Mode), the gateway sets appConfig.demoMode to \"live\". "
        "This causes every downstream operation to use real Azure OpenAI calls, real Foundry "
        "API registrations, and live MCP server health probes. When set to \"simulated\", "
        "all operations use hardcoded mock data, timed animations, and synthetic results."
    )
    doc.add_paragraph(
        "This document focuses exclusively on the Real Mode code paths."
    )

    doc.add_page_break()

    # ========================================================
    # 3. TEST TAB
    # ========================================================
    doc.add_heading("3. Test Tab - Backend Walkthrough", level=1)
    doc.add_paragraph(
        "The Test tab validates that a designed workflow is ready for deployment by running "
        "structured test scenarios against the live infrastructure. It does not process actual "
        "contracts; instead it probes design completeness, tool coverage, and server connectivity."
    )

    doc.add_heading("3.1 Tab Initialisation", level=2)
    doc.add_paragraph(
        "When the user switches to the Test tab, the UI calls syncTestTab() which triggers "
        "two parallel API requests:"
    )
    add_styled_table(doc,
        ["API Call", "Endpoint", "Purpose"],
        [
            ["loadTestScenarios()", "GET /api/v1/test-scenarios", "Fetches the array of predefined test scenarios from data/test-scenarios.json"],
            ["loadToolRegistry()", "GET /api/v1/tools", "Fetches the list of all MCP servers with their online/offline status and registered tool names"],
        ],
        col_widths=[4, 5, 8],
    )

    doc.add_heading("3.2 Test Scenarios (data/test-scenarios.json)", level=2)
    doc.add_paragraph(
        "The system ships with predefined test scenarios. Each scenario is a JSON object with:"
    )
    doc.add_paragraph("id - A unique scenario identifier (e.g., \"high-risk-msa\")", style="List Bullet")
    doc.add_paragraph("name - Human-readable name (e.g., \"High-Risk Master Service Agreement\")", style="List Bullet")
    doc.add_paragraph("description - What the scenario tests", style="List Bullet")
    doc.add_paragraph(
        "requiredCapabilities - Array of capability IDs the workflow must cover (e.g., "
        "[\"intake\", \"drafting\", \"review\", \"compliance\", \"negotiation\", \"approval\"])",
        style="List Bullet"
    )
    doc.add_paragraph("requiresHumanReview - Boolean flag for Human-In-The-Loop expectation", style="List Bullet")
    doc.add_paragraph("prefersParallel - Boolean flag for parallel stage recommendation", style="List Bullet")
    doc.add_paragraph(
        "expectations - Array of human-readable expectations (e.g., \"Intake should classify "
        "the agreement as high-risk\")",
        style="List Bullet"
    )

    doc.add_heading("3.3 Running a Test Scenario (Real Mode)", level=2)
    doc.add_paragraph(
        "When the user selects a scenario and clicks \"Run Scenario\", the UI calls "
        "evaluateWorkflowScenario(workflow, scenario). This function executes a multi-step "
        "validation process entirely in the browser, but with live backend health probes "
        "in Real Mode. The steps are:"
    )

    doc.add_heading("Step 1: Design Validation", level=3)
    doc.add_paragraph(
        "The function calls getWorkflowValidation(workflow) which checks the structural "
        "integrity of the workflow definition: are agents defined? Do agents have tool "
        "assignments? Are stage mappings present? Returns an object with errors, warnings, "
        "findings, and an isValid boolean. A single assertion is added to the result."
    )

    doc.add_heading("Step 2: Stage Extraction", level=3)
    doc.add_paragraph(
        "getWorkflowStagesForTesting(workflow) maps the workflow's contract_stage_map to "
        "extract an ordered list of stages, each associated with the agents assigned to it. "
        "This is used for coverage analysis and the stage trace visualisation."
    )

    doc.add_heading("Step 3: Tool Coverage Check (Real Mode)", level=3)
    doc.add_paragraph(
        "For each capability ID in scenario.requiredCapabilities, the system maps the "
        "capability to specific MCP tool names using a CAPABILITY_RULES lookup table. "
        "For example, \"compliance\" maps to tools [\"check_policy\", \"flag_risk\"]. "
        "In Real Mode, this step queries the live tool registry:"
    )
    doc.add_paragraph(
        "1. Calls getToolRegistryStatus(capabilityId) which checks the cached /api/v1/tools "
        "response to determine if the relevant MCP server is online.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. If the server IS online, the assertion passes with detail: \"[Live] {capability} "
        "online on {server_name} ({tool1}, {tool2}).\"",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. If the server IS offline, the assertion fails with detail: \"[Live] MCP server "
        "offline. {capability} tools unavailable.\"",
        style="List Bullet"
    )
    doc.add_paragraph(
        "This is the key difference from Simulated Mode, which only checks whether the "
        "workflow definition includes an agent with the right keywords, without verifying "
        "actual server availability."
    )

    doc.add_heading("Step 4: Human Checkpoint Check", level=3)
    doc.add_paragraph(
        "If scenario.requiresHumanReview is true, the system checks whether the workflow "
        "includes at least one agent with kind = \"human\". If Human-In-The-Loop is expected "
        "but not configured, the assertion fails."
    )

    doc.add_heading("Step 5: Parallel Review Check", level=3)
    doc.add_paragraph(
        "If scenario.prefersParallel is true, the system checks whether any extracted stage "
        "has more than one agent assigned (indicating parallel processing). If no parallel "
        "stage is found, a warning assertion is added."
    )

    doc.add_heading("Step 6: Gateway Connectivity Check (Real Mode Only)", level=3)
    doc.add_paragraph(
        "This step is exclusive to Real Mode. The UI fetches GET /api/v1/health from the "
        "gateway. The gateway responds with:"
    )
    doc.add_paragraph(
        '{ "mode": "live", "servers": { "contract-intake-mcp": "online", '
        '"contract-extraction-mcp": "online", ... } }'
    )
    doc.add_paragraph(
        "The test counts online vs total servers. If all are online, the assertion passes. "
        "If some are online, a warning is issued. If none are online, it fails. The detail "
        "reads: \"[Live] {online}/{total} MCP servers online\"."
    )

    doc.add_heading("3.4 Test Result Structure", level=2)
    doc.add_paragraph("The evaluateWorkflowScenario function returns an object containing:")
    add_styled_table(doc,
        ["Field", "Type", "Description"],
        [
            ["scenario", "Object", "The selected scenario definition"],
            ["workflow", "Object", "The active workflow definition"],
            ["stages", "Array", "Extracted stages with agent mappings"],
            ["assertions", "Array", "Each assertion has: status (pass/warn/fail), label, detail, timestamp"],
            ["passCount", "Number", "Count of passed assertions"],
            ["warnCount", "Number", "Count of warning assertions"],
            ["failCount", "Number", "Count of failed assertions"],
            ["verdict", "String", "Overall verdict: pass, warn, or fail"],
        ],
        col_widths=[3, 2, 12],
    )

    doc.add_heading("3.5 UI Rendering", level=2)
    doc.add_paragraph(
        "The result is rendered as a summary badge ([PASS]/[WARN]/[FAIL] with counts) "
        "followed by a list of individual assertions with colour-coded status badges. "
        "A stage trace diagram shows the visual flow of agents through each pipeline stage."
    )

    doc.add_page_break()

    # ========================================================
    # 4. DEPLOY TAB
    # ========================================================
    doc.add_heading("4. Deploy Tab - Backend Walkthrough", level=1)
    doc.add_paragraph(
        "The Deploy tab registers AI agents on Azure AI Foundry, verifies model access, "
        "validates content safety filters, runs evaluation prompts against each registered "
        "agent, and performs health checks. This is a full deployment pipeline."
    )

    doc.add_heading("4.1 Triggering the Deploy Pipeline", level=2)
    doc.add_paragraph(
        "When the user clicks \"Deploy Pipeline\", the UI calls runDeployPipelineReal() "
        "which sends:"
    )
    doc.add_paragraph(
        "POST /api/v1/deploy/pipeline\n"
        "Header: x-admin-key: {DEPLOY_ADMIN_KEY}\n"
        "Timeout: 180 seconds (3 minutes)",
        style="List Bullet"
    )

    doc.add_heading("4.2 Security Gate", level=2)
    doc.add_paragraph(
        "The deploy route handler calls ensureDeployAdminAccess() which enforces "
        "authentication in Real Mode:"
    )
    doc.add_paragraph(
        "1. If DEPLOY_ADMIN_KEY is not configured on the server, returns 503 "
        "(Service Unavailable) with error \"deploy_admin_not_configured\".",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. If the x-admin-key header does not match the configured key, returns 401 "
        "(Unauthorized).",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. Only on a valid key match does the pipeline proceed.",
        style="List Bullet"
    )

    doc.add_heading("4.3 Deployment Pipeline Stages (gateway/src/services/foundryDeploy.ts)", level=2)
    doc.add_paragraph(
        "The gateway calls deployToFoundry(config, onProgress) which executes six sequential "
        "stages. Each stage returns a StageResult with: name, status (pending/running/passed/"
        "failed/skipped), duration_ms, optional details, and optional error. If any critical "
        "stage fails, the pipeline short-circuits and returns immediately."
    )

    doc.add_heading("Stage 1: Preflight", level=3)
    doc.add_paragraph(
        "Purpose: Verify that the Azure AI Foundry endpoint is reachable and the API key "
        "has valid data-plane access."
    )
    doc.add_paragraph("Action:", style="List Bullet")
    doc.add_paragraph(
        "Sends GET /openai/models?api-version=2024-10-21 to the Foundry endpoint using "
        "the configured authentication (API key or Managed Identity). This uses the data-plane "
        "models endpoint rather than the management-plane deployments endpoint because API keys "
        "only have data-plane access."
    )
    doc.add_paragraph("Success: Returns the number of models found on the endpoint.", style="List Bullet")
    doc.add_paragraph("Failure: Returns \"API access denied\" with the HTTP status code.", style="List Bullet")
    doc.add_paragraph(
        "Short-circuit: If preflight fails, the entire pipeline stops immediately.",
        style="List Bullet"
    )

    doc.add_heading("Stage 2: Model Verification", level=3)
    doc.add_paragraph(
        "Purpose: Confirm that the specific model deployment (e.g., gpt-4o) is accessible and "
        "can serve completions."
    )
    doc.add_paragraph("Action:", style="List Bullet")
    doc.add_paragraph(
        "Sends a minimal chat completion request to "
        "POST /openai/deployments/{model}/chat/completions?api-version=2024-10-21 "
        "with a single \"ping\" message and max_tokens=1. This verifies the deployment "
        "exists and is functional."
    )
    doc.add_paragraph(
        "Success: Returns deployment name, model identifier, and \"verified_via: chat_completion\".",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Failure (404): \"Model deployment '{model}' was not found. Provision it through "
        "the Azure deployment workflow before running agent registration.\"",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Short-circuit: If model verification fails, the pipeline stops.",
        style="List Bullet"
    )

    doc.add_heading("Stage 3: Agent Registration (Idempotent)", level=3)
    doc.add_paragraph(
        "Purpose: Create (or reuse) OpenAI Assistant objects on Azure AI Foundry for each "
        "of the six pipeline agents."
    )
    doc.add_paragraph("The system defines six agents:")
    add_styled_table(doc,
        ["Agent", "System Prompt File", "MCP Tools"],
        [
            ["Contract Intake Agent", "intake-system.md", "upload_contract, classify_document, extract_metadata"],
            ["Contract Drafting Agent", "drafting-system.md", "extract_clauses, identify_parties, extract_dates_values"],
            ["Contract Internal Review Agent", "review-system.md", "get_audit_log, create_audit_entry"],
            ["Contract Compliance Agent", "compliance-system.md", "check_policy, flag_risk, get_policy_rules"],
            ["Contract Negotiation Agent", "negotiation-system.md", "route_approval, notify_stakeholder"],
            ["Contract Approval Agent", "approval-system.md", "route_approval, escalate_to_human, notify_stakeholder"],
        ],
        col_widths=[5, 4, 8],
    )

    doc.add_paragraph()
    doc.add_paragraph("The registration process for each agent is:")
    doc.add_paragraph(
        "1. List existing assistants: GET /openai/assistants?api-version=2024-05-01-preview&limit=100 "
        "on the project endpoint. Filters by metadata.domain = \"contract-management\".",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. If an assistant with the same name already exists, it is reused (idempotent). "
        "The existing foundry_agent_id is returned without creating a duplicate.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. If no existing match is found, the system loads the system prompt from the "
        "prompts/ directory (e.g., prompts/intake-system.md), builds the tools array as "
        "OpenAI function-calling tool definitions, and sends POST /openai/assistants with: "
        "model, name, description, instructions (system prompt), tools, temperature (0.1), "
        "and metadata tags (domain, pipeline_role, mcp_tools, version).",
        style="List Bullet"
    )
    doc.add_paragraph(
        "4. The response contains the assistant ID (e.g., asst_abc123) which becomes the "
        "foundry_agent_id used for subsequent evaluation and health checks.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Result details include: registered count, total count, reused count, newly created "
        "count, and total tool definitions registered.",
        style="List Bullet"
    )

    doc.add_heading("Stage 4: Content Safety Verification", level=3)
    doc.add_paragraph(
        "Purpose: Verify that Azure content safety filters are active on the model deployment."
    )
    doc.add_paragraph("Action:", style="List Bullet")
    doc.add_paragraph(
        "Sends a test chat completion with a benign probe message (\"Test: verify content "
        "safety filters are active.\") and inspects the response for content_filter_results "
        "in the choices array."
    )
    doc.add_paragraph(
        "Scenario A - Filters detected in response: The stage passes, confirming filters "
        "for categories: hate, sexual, violence, self_harm, jailbreak.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Scenario B - 400 error with \"content_filter\" in body: This actually means filters "
        "ARE active (they triggered on the test). The stage passes.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Scenario C - No filters detected: The system attempts to activate them by reading "
        "the current deployment config (GET /openai/deployments/{model}) and updating it "
        "(PUT) with contentFilter: { defaultPolicy: true }. If this fails (often due to "
        "API limitations), the stage reports that filters must be configured manually via "
        "the Azure Portal.",
        style="List Bullet"
    )

    doc.add_heading("Stage 5: Agent Evaluation", level=3)
    doc.add_paragraph(
        "Purpose: Run a quick smoke test against each registered agent to verify it can "
        "process its domain-specific prompt and produce a response."
    )
    doc.add_paragraph("For each registered agent, the system:")
    doc.add_paragraph(
        "1. Creates a Thread: POST /openai/threads on the project endpoint.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. Posts an evaluation message: POST /openai/threads/{thread_id}/messages with "
        "the agent-specific evalPrompt (e.g., for the Compliance Agent: \"Assess this clause "
        "for policy risk: 'Vendor liability is capped at $10,000,000 and personal data may be "
        "transferred outside approved jurisdictions.'\").",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. Starts a Run: POST /openai/threads/{thread_id}/runs with the assistant_id.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "4. Polls the run status every 2 seconds (up to 15 polls = 30 seconds max) until "
        "it reaches \"completed\", \"failed\", or \"cancelled\".",
        style="List Bullet"
    )
    doc.add_paragraph(
        "5. Cleans up the thread: DELETE /openai/threads/{thread_id}.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Result: Reports test_count, passed, accuracy percentage, and any failure details.",
        style="List Bullet"
    )

    doc.add_heading("Stage 6: Health Check", level=3)
    doc.add_paragraph(
        "Purpose: Verify all registered agents are still accessible on the Foundry API."
    )
    doc.add_paragraph(
        "For each registered agent, sends GET /openai/assistants/{foundry_agent_id} on the "
        "project endpoint. Counts how many respond successfully. If all are healthy, the stage "
        "passes."
    )

    doc.add_heading("4.4 Deploy Pipeline Result", level=2)
    doc.add_paragraph(
        "The complete DeployPipelineResult returned to the UI contains:"
    )
    add_styled_table(doc,
        ["Field", "Type", "Description"],
        [
            ["pipeline_id", "String", "Unique deployment ID (e.g., deploy-a1b2c3d4)"],
            ["mode", "String", "\"live\" for Real Mode"],
            ["stages", "Array<StageResult>", "Results for all 6 stages (Preflight, Model Deployment, Agent Registration, Content Safety, Evaluation, Health Check)"],
            ["agents", "Array<FoundryAgentInfo>", "Each agent: name, foundry_agent_id, model, status (registered/failed), tools_count"],
            ["security.identity_access", "Array", "API Key auth, RBAC roles, Data residency checks"],
            ["security.content_safety", "Array", "Content filters, Jailbreak protection, PII redaction checks"],
            ["evaluation", "Object", "test_count, passed, accuracy percentage"],
            ["summary", "Object", "agents_deployed, tools_registered, errors, total_duration_ms"],
        ],
        col_widths=[4, 3.5, 9.5],
    )

    doc.add_heading("4.5 UI Rendering", level=2)
    doc.add_paragraph(
        "The UI maps the six backend stages to four visual stage elements: Preflight maps "
        "to the Build stage, Model Deployment maps to the Test stage, Agent Registration "
        "maps to the Deploy stage, and Health Check maps to the Register stage. Each element "
        "displays a [PASS]/[FAIL] badge and the actual duration. The agent registry table "
        "is populated with each agent's name, Foundry ID, status badge, and tool count."
    )

    doc.add_page_break()

    # ========================================================
    # 5. LIVE TAB
    # ========================================================
    doc.add_heading("5. Live Tab - Backend Walkthrough", level=1)
    doc.add_paragraph(
        "The Live tab is the operational heart of the platform. It processes actual contracts "
        "through the full six-agent AI pipeline, streaming results in real-time via WebSocket. "
        "This is where Azure OpenAI calls happen, clauses are extracted, compliance is assessed, "
        "and approval decisions are routed."
    )

    doc.add_heading("5.1 Contract Submission (Phase 1)", level=2)
    doc.add_paragraph(
        "The user selects a sample contract from a dropdown or uploads a contract file via "
        "drag-and-drop. The UI sends:"
    )
    doc.add_paragraph(
        'POST /api/v1/contracts\n'
        '{ "text": "<full contract text>", "filename": "nda-acme-beta.txt" }',
        style="List Bullet"
    )
    doc.add_paragraph(
        "Validation: The gateway validates the contract text is present and does not exceed "
        "50,000 characters. Invalid requests receive a 400 error."
    )
    doc.add_paragraph(
        "The gateway returns 202 Accepted immediately with:"
    )
    doc.add_paragraph(
        '{ "contract_id": "contract-a1b2c3d4", "status": "processing", '
        '"message": "Contract submitted. Follow /ws/workflow for real-time updates." }',
        style="List Bullet"
    )
    doc.add_paragraph(
        "This is a non-blocking response. The actual pipeline runs in the background."
    )

    doc.add_heading("5.2 WebSocket Connection", level=2)
    doc.add_paragraph(
        "The UI maintains a persistent WebSocket connection to WS /ws/workflow. All pipeline "
        "progress events are broadcast through this channel. Events use a typed structure with "
        "event (event type), contract_id, status, optional result payload, latency_ms, and "
        "timestamp."
    )

    doc.add_heading("5.3 Background Pipeline Execution (Phase 2)", level=2)
    doc.add_paragraph(
        "The gateway calls runPipeline(contractText, filename, adapter, broadcast, contractId) "
        "from gateway/src/orchestrator/pipeline.ts. This function orchestrates six sequential "
        "AI agent stages with retry logic, validation, audit logging, and real-time broadcasting."
    )

    doc.add_heading("LLM Adapter", level=3)
    doc.add_paragraph(
        "The gateway creates an AzureOpenAIAdapter (implementing the ILlmAdapter interface) "
        "that sends requests to the configured Azure OpenAI endpoint. A TrackingAdapter wraps "
        "this to capture token usage (tokens_in, tokens_out) for each call. The adapter is "
        "reset between stages to isolate token metrics."
    )

    doc.add_heading("Retry Logic", level=3)
    doc.add_paragraph(
        "Every agent call is wrapped in withRetry() which implements exponential backoff:"
    )
    doc.add_paragraph("Max retries: 3 attempts (4 total including initial try)", style="List Bullet")
    doc.add_paragraph("Base delay: 1,000ms, with exponential growth (1s, 2s, 4s)", style="List Bullet")
    doc.add_paragraph("Max delay cap: 10,000ms", style="List Bullet")
    doc.add_paragraph(
        "Each retry logs an audit entry: \"Attempt {n}/{total} failed: {error}. Retrying in {delay}ms\"",
        style="List Bullet"
    )
    doc.add_paragraph(
        "If all retries are exhausted, a PipelineError is thrown with the stage name, agent "
        "name, original error, and contract ID.",
        style="List Bullet"
    )

    doc.add_heading("5.4 Stage 1: Intake Agent", level=2)
    doc.add_paragraph("Source: agents/src/intakeAgent.ts")
    doc.add_paragraph("Purpose: Classify the contract type and extract high-level metadata.")
    doc.add_paragraph("Process:")
    doc.add_paragraph(
        "1. Loads the system prompt from prompts/intake-system.md which contains detailed "
        "instructions for contract classification.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. Sends a chat completion request to Azure OpenAI with the system prompt and the "
        "full contract text, requesting JSON response format.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. Parses the JSON response to extract: contract_type (e.g., \"NDA\", \"MSA\"), "
        "confidence_score (0-1), parties (array of organisation names), and metadata "
        "(title, effective_date, expiry_date, value, currency, jurisdiction).",
        style="List Bullet"
    )
    doc.add_paragraph(
        "4. Validation: If type is missing or confidence < 0.3, the pipeline throws an error.",
        style="List Bullet"
    )
    doc.add_paragraph("Post-processing:", style="List Bullet")
    doc.add_paragraph(
        "- Stores a trace entry with agent name, tool (classify_document), input/output, "
        "latency, and token counts.\n"
        "- Updates the contract record in contractStore with the classified type and confidence.\n"
        "- Writes an audit log entry: \"Classified as {type} with {confidence} confidence\".\n"
        "- Broadcasts an agent_step_complete WebSocket event with the full result."
    )

    doc.add_heading("5.5 Stage 2: Extraction Agent", level=2)
    doc.add_paragraph("Source: agents/src/extractionAgent.ts")
    doc.add_paragraph("Purpose: Extract individual clauses, parties, dates, and monetary values.")
    doc.add_paragraph("Process:")
    doc.add_paragraph(
        "1. Loads prompts/drafting-system.md (the extraction/drafting system prompt).",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. Sends the contract text to Azure OpenAI requesting structured clause extraction.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. Returns an array of extracted clauses, each with clause text, type, risk "
        "indicators, and positions.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "4. Validation: If no clauses are extracted (empty array), the pipeline throws an error.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Post-processing: Trace stored, audit entry (\"Extracted {n} clauses\"), WebSocket broadcast.",
        style="List Bullet"
    )

    doc.add_heading("5.6 Stage 3: Internal Review Agent", level=2)
    doc.add_paragraph("Source: agents/src/reviewAgent.ts")
    doc.add_paragraph("Purpose: Analyse extracted clauses for material changes and unresolved items.")
    doc.add_paragraph("Process:")
    doc.add_paragraph(
        "1. Receives the extracted clauses (not the raw text) from Stage 2.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. Sends them to Azure OpenAI with the review system prompt.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. Returns: materialChanges (array of flagged modifications) and unresolvedItems "
        "(issues requiring attention).",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Audit entry: \"Review: {n} material changes, {m} unresolved items\"",
        style="List Bullet"
    )

    doc.add_heading("5.7 Stage 4: Compliance Agent", level=2)
    doc.add_paragraph("Source: agents/src/complianceAgent.ts")
    doc.add_paragraph("Purpose: Assess clauses against organisational policies and flag risk.")
    doc.add_paragraph("Process:")
    doc.add_paragraph(
        "1. Receives the extracted clauses from Stage 2.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. Sends them to Azure OpenAI with the compliance system prompt which contains "
        "policy rules and risk assessment criteria.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. Returns: overallRisk (e.g., \"HIGH\", \"MEDIUM\", \"LOW\"), flagsCount (number "
        "of policy violations), and detailed flag descriptions.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Validation: If overallRisk is missing, the pipeline throws an error.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Audit entry: \"Risk: {level}, Flags: {count}\"",
        style="List Bullet"
    )

    doc.add_heading("5.8 Stage 5: Negotiation Agent", level=2)
    doc.add_paragraph("Source: agents/src/negotiationAgent.ts")
    doc.add_paragraph("Purpose: Assess counterparty positions and recommend negotiation strategy.")
    doc.add_paragraph("Process:")
    doc.add_paragraph(
        "1. Receives: extracted clauses, overall risk level, and flags count from previous stages.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. Sends to Azure OpenAI with the negotiation system prompt.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. Returns: counterpartyPositions (array of negotiation points) and "
        "escalationRequired (boolean).",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Audit entry: \"Negotiation: {n} positions, escalation={true/false}\"",
        style="List Bullet"
    )

    doc.add_heading("5.9 Stage 6: Approval Agent", level=2)
    doc.add_paragraph("Source: agents/src/approvalAgent.ts")
    doc.add_paragraph("Purpose: Make the final routing decision - auto-approve or escalate to a human.")
    doc.add_paragraph("Process:")
    doc.add_paragraph(
        "1. Receives: overall risk level and flags count.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. Sends to Azure OpenAI with the approval system prompt.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. Returns: action (\"auto_approve\" or \"escalate_to_human\") and reasoning (text).",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Validation: If action is not one of the two expected values, the pipeline throws an error.",
        style="List Bullet"
    )

    doc.add_heading("5.10 Approval Routing", level=2)
    doc.add_paragraph("Based on the Approval Agent's decision, one of two paths is taken:")
    doc.add_paragraph(
        "Auto-Approve Path: The contract status is set to \"approved\", completed_at is "
        "timestamped, an audit entry with action \"approved\" is stored, and a pipeline_status "
        "event with status \"approved\" is broadcast. The pipeline completes.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Escalate Path: The contract status is set to \"awaiting_review\", an audit entry "
        "with action \"escalated\" is stored, and an agent_step_complete event with status "
        "\"awaiting_human_review\" is broadcast. The UI displays the Human-In-The-Loop panel.",
        style="List Bullet"
    )

    doc.add_heading("5.11 Human-In-The-Loop (HITL) Review", level=2)
    doc.add_paragraph(
        "When the Approval Agent escalates, the HITL panel appears with three buttons: "
        "Approve, Reject, and Request Changes, plus a comment field."
    )
    doc.add_paragraph("When the human reviewer submits a decision, the UI sends:")
    doc.add_paragraph(
        'POST /api/v1/contracts/{contract_id}/review\n'
        '{ "decision": "approve|reject|request_changes", "comment": "...", "reviewer": "human" }',
        style="List Bullet"
    )
    doc.add_paragraph("The gateway handler:")
    doc.add_paragraph(
        "1. Validates the decision is one of: approve, reject, request_changes.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "2. Maps the decision to a contract status: approve -> \"approved\", "
        "reject -> \"rejected\", request_changes -> \"awaiting_review\".",
        style="List Bullet"
    )
    doc.add_paragraph(
        "3. Updates the contract record in contractStore with the new status and completed_at "
        "timestamp.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "4. Stores an audit entry with agent=\"human\", the decision action, and the "
        "reviewer's comment as reasoning.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "5. Broadcasts a pipeline_status WebSocket event with the final status.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "6. Returns the decision confirmation to the UI.",
        style="List Bullet"
    )

    doc.add_heading("5.12 Trace Storage", level=2)
    doc.add_paragraph(
        "After the pipeline completes (regardless of approval path), all trace entries "
        "are persisted via storeTraces(contractId, traces). Each trace captures: unique ID, "
        "contract_id, agent name, tool name, input summary, full output, latency_ms, "
        "tokens_in, tokens_out, and timestamp. This data feeds into the evaluation and "
        "drift detection systems."
    )

    doc.add_heading("5.13 Error Handling", level=2)
    doc.add_paragraph(
        "If any stage fails after all retry attempts are exhausted, the pipeline catch block:"
    )
    doc.add_paragraph(
        "- Updates the contract status to \"failed\" in contractStore.\n"
        "- Broadcasts a pipeline_status event with status \"error\" and the error message.\n"
        "- The UI displays the error state on the relevant workflow node."
    )

    doc.add_page_break()

    # ========================================================
    # 6. API REFERENCE
    # ========================================================
    doc.add_heading("6. API Reference Summary", level=1)
    doc.add_paragraph(
        "The following table summarises all key API endpoints used by the three tabs "
        "in Real Mode."
    )

    add_styled_table(doc,
        ["Endpoint", "Method", "Used By", "Purpose"],
        [
            ["/api/v1/mode", "POST", "Mode Switch", "Toggle between live and simulated mode"],
            ["/api/v1/health", "GET", "Test, Live", "Get MCP server online/offline status and current mode"],
            ["/api/v1/test-scenarios", "GET", "Test", "Load predefined test scenario definitions"],
            ["/api/v1/tools", "GET", "Test", "List all MCP servers with tool names and live status"],
            ["/api/v1/deploy/pipeline", "POST", "Deploy", "Execute the full Foundry deployment pipeline (requires x-admin-key)"],
            ["/api/v1/contracts", "POST", "Live", "Submit a contract for processing (returns 202 Accepted)"],
            ["/api/v1/contracts", "GET", "Live", "List all processed contracts"],
            ["/api/v1/contracts/:id", "GET", "Live", "Get contract details and current status"],
            ["/api/v1/contracts/:id/review", "POST", "Live", "Submit HITL review decision (approve/reject/request_changes)"],
            ["/api/v1/sample-contracts", "GET", "Live", "List available sample contracts"],
            ["/api/v1/sample-contracts/:filename", "GET", "Live", "Load a sample contract's text"],
            ["/api/v1/workflows", "GET", "All", "List, save, or activate workflow definitions"],
            ["WS /ws/workflow", "WebSocket", "Live", "Real-time pipeline event stream"],
        ],
        col_widths=[5, 1.5, 2, 8.5],
    )

    doc.add_page_break()

    # ========================================================
    # 7. DATA FLOW DIAGRAMS
    # ========================================================
    doc.add_heading("7. Data Flow Diagrams", level=1)

    doc.add_heading("7.1 Test Tab Data Flow", level=2)
    doc.add_paragraph(
        "User clicks Test tab\n"
        "  |-> GET /api/v1/test-scenarios -> Load scenario definitions\n"
        "  |-> GET /api/v1/tools -> Load MCP server status + tool lists\n"
        "  |\n"
        "User selects scenario, clicks \"Run Scenario\"\n"
        "  |-> evaluateWorkflowScenario() [client-side]\n"
        "  |   |-> Step 1: getWorkflowValidation(workflow) -> design check\n"
        "  |   |-> Step 2: getWorkflowStagesForTesting(workflow) -> stage map\n"
        "  |   |-> Step 3: For each capability -> getToolRegistryStatus() -> [Live] server check\n"
        "  |   |-> Step 4: Human checkpoint check\n"
        "  |   |-> Step 5: Parallel stage check\n"
        "  |   |-> Step 6: GET /api/v1/health -> [Live] gateway connectivity\n"
        "  |-> Render assertions with pass/warn/fail badges"
    )

    doc.add_heading("7.2 Deploy Tab Data Flow", level=2)
    doc.add_paragraph(
        "User clicks \"Deploy Pipeline\"\n"
        "  |-> POST /api/v1/deploy/pipeline (x-admin-key header)\n"
        "  |   |-> ensureDeployAdminAccess() -> auth check\n"
        "  |   |-> deployToFoundry(config)\n"
        "  |       |-> Stage 1: Preflight  -> GET /openai/models (verify API access)\n"
        "  |       |-> Stage 2: Model Verification -> POST chat/completions (ping)\n"
        "  |       |-> Stage 3: Agent Registration -> POST /openai/assistants x6\n"
        "  |       |   |-> For each agent: load prompt, build tools, create assistant\n"
        "  |       |   |-> Idempotent: reuse existing if name matches\n"
        "  |       |-> Stage 4: Content Safety -> POST chat/completions (probe)\n"
        "  |       |-> Stage 5: Evaluation -> Thread + Message + Run per agent\n"
        "  |       |-> Stage 6: Health Check -> GET /openai/assistants/{id} per agent\n"
        "  |-> Return DeployPipelineResult\n"
        "  |-> UI maps stages to visual elements, populates agent registry table"
    )

    doc.add_heading("7.3 Live Tab Data Flow", level=2)
    doc.add_paragraph(
        "User selects/uploads contract\n"
        "  |-> POST /api/v1/contracts { text, filename }\n"
        "  |   |-> Returns 202 { contract_id, status: \"processing\" }\n"
        "  |   |-> Background: runPipeline()\n"
        "  |       |-> Broadcast: pipeline_status = processing_started\n"
        "  |       |\n"
        "  |       |-> Stage 1: runIntakeAgent()\n"
        "  |       |   |-> Load intake-system.md prompt\n"
        "  |       |   |-> Azure OpenAI chat completion (JSON response)\n"
        "  |       |   |-> Parse: type, confidence, parties, metadata\n"
        "  |       |   |-> Validate (confidence >= 0.3)\n"
        "  |       |   |-> Store trace, update contract, log audit\n"
        "  |       |   |-> Broadcast: agent_step_complete (intake)\n"
        "  |       |\n"
        "  |       |-> Stage 2: runExtractionAgent()\n"
        "  |       |   |-> Azure OpenAI -> extract clauses array\n"
        "  |       |   |-> Validate (clauses.length > 0)\n"
        "  |       |   |-> Broadcast: agent_step_complete (extraction)\n"
        "  |       |\n"
        "  |       |-> Stage 3: runReviewAgent(clauses)\n"
        "  |       |   |-> Azure OpenAI -> materialChanges, unresolvedItems\n"
        "  |       |   |-> Broadcast: agent_step_complete (review)\n"
        "  |       |\n"
        "  |       |-> Stage 4: runComplianceAgent(clauses)\n"
        "  |       |   |-> Azure OpenAI -> overallRisk, flagsCount\n"
        "  |       |   |-> Validate (overallRisk present)\n"
        "  |       |   |-> Broadcast: agent_step_complete (compliance)\n"
        "  |       |\n"
        "  |       |-> Stage 5: runNegotiationAgent(clauses, risk, flags)\n"
        "  |       |   |-> Azure OpenAI -> counterpartyPositions, escalationRequired\n"
        "  |       |   |-> Broadcast: agent_step_complete (negotiation)\n"
        "  |       |\n"
        "  |       |-> Stage 6: runApprovalAgent(risk, flags)\n"
        "  |       |   |-> Azure OpenAI -> action, reasoning\n"
        "  |       |   |-> Validate (action = auto_approve | escalate_to_human)\n"
        "  |       |   |\n"
        "  |       |   |-> IF auto_approve:\n"
        "  |       |   |   |-> status = approved, audit = approved\n"
        "  |       |   |   |-> Broadcast: pipeline_status = approved\n"
        "  |       |   |\n"
        "  |       |   |-> IF escalate_to_human:\n"
        "  |       |       |-> status = awaiting_review, audit = escalated\n"
        "  |       |       |-> Broadcast: awaiting_human_review\n"
        "  |       |       |-> UI shows HITL panel\n"
        "  |       |       |-> Human clicks Approve/Reject/Request Changes\n"
        "  |       |       |-> POST /api/v1/contracts/{id}/review\n"
        "  |       |       |-> Update contract, audit, broadcast final status\n"
        "  |       |\n"
        "  |       |-> storeTraces(contractId, traces)\n"
        "  |\n"
        "  |-> UI WebSocket receives all events, updates:\n"
        "      |-> Workflow node states (processing -> complete)\n"
        "      |-> Progress bars and latency displays\n"
        "      |-> Activity log entries\n"
        "      |-> Contract detail panel\n"
        "      |-> HITL panel visibility"
    )

    doc.add_page_break()

    # ========================================================
    # APPENDIX
    # ========================================================
    doc.add_heading("Appendix: Key Source Files", level=1)
    add_styled_table(doc,
        ["File", "Purpose"],
        [
            ["start.ts", "Application entry point; spawns MCP servers and gateway"],
            ["gateway/src/routes/contracts.ts", "Contract submission and HITL review endpoints"],
            ["gateway/src/routes/deploy.ts", "Deploy pipeline endpoint with admin auth gate"],
            ["gateway/src/orchestrator/pipeline.ts", "Six-stage pipeline orchestrator with retry logic"],
            ["gateway/src/services/foundryDeploy.ts", "Foundry deployment: preflight, model check, agent registration, safety, eval, health"],
            ["gateway/src/adapters/trackingAdapter.ts", "LLM adapter wrapper that tracks token usage per call"],
            ["gateway/src/stores/contractStore.ts", "In-memory contract and audit data stores"],
            ["agents/src/intakeAgent.ts", "Intake Agent: contract classification and metadata extraction"],
            ["agents/src/extractionAgent.ts", "Extraction Agent: clause extraction from contract text"],
            ["agents/src/reviewAgent.ts", "Review Agent: internal review for material changes"],
            ["agents/src/complianceAgent.ts", "Compliance Agent: policy checking and risk flagging"],
            ["agents/src/negotiationAgent.ts", "Negotiation Agent: counterparty position assessment"],
            ["agents/src/approvalAgent.ts", "Approval Agent: auto-approve vs escalate decision"],
            ["agents/src/agentConfig.ts", "Agent definitions and system prompt loader"],
            ["ui/app.js", "Frontend: tab management, test runner, deploy/live mode logic"],
            ["ui/api.js", "API client: apiCall(), WebSocket connection management"],
            ["data/test-scenarios.json", "Predefined test scenario definitions"],
            ["prompts/*.md", "System prompts for each agent (intake, drafting, review, compliance, negotiation, approval)"],
        ],
        col_widths=[6, 11],
    )

    # Save
    output_path = os.path.join(os.path.dirname(__file__), "docs",
                               "AgentOps_Backend_RealMode_Test_Deploy_Live.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
