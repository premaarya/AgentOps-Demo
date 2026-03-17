from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('The Complete, Simple Guide: How Our AgentOps Application Runs in the Cloud', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Introduction
    doc.add_heading('1. The Big Picture: What is "Deployment"?', level=1)
    doc.add_paragraph('Think of building an application like running a restaurant. When developers are first writing the code, they are cooking in their own private home kitchen (their laptop). They can taste the food, adjust the recipes, and make sure everything is perfect, but the general public cannot walk into their house to eat.')
    doc.add_paragraph('To let the world use our application, we need to rent a professional, public space, hire staff to keep the lights on, and set up a system that can handle hundreds of customers at once. This process of moving from the "home kitchen" to the "public restaurant" is what we call Deployment.')
    
    # Infrastructure
    doc.add_heading('2. Where Does Our App Live? (Our Landlord)', level=1)
    doc.add_paragraph('Instead of buying our own physical computers and plugging them in to the wall, we rent space from Microsoft Azure. Azure acts as our landlord and property manager.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Azure App Service: ').bold = True
    p.add_run('This is the specific "office space" we rent from Microsoft. We just hand them our application code, and Microsoft guarantees that the computer running it will stay powered on, connected to fast internet, and physically secure. If the physical computer breaks, Microsoft instantly moves our code to a new one without us even noticing.')
    
    # Architecture
    doc.add_heading('3. The Team: How Our App Works Inside', level=1)
    doc.add_paragraph('Our application is not just one big block of code; it is a team of different workers communicating with each other:')
    
    parts = [
        ('The Dashboard (Frontend)', 'This is the visual website you see on your screen. It is like the dining room and the menu, where users click buttons and read text.'),
        ('The API Gateway (The Waiter)', 'When you click a button on the dashboard, the Gateway takes your request and figures out which specialized chef needs to handle it.'),
        ('The Agents / MCP Servers (The Specialized Chefs)', 'We have 8 different specialized mini-programs running in the background. One is an expert at reading documents (Extraction), another is an expert at checking rules (Compliance), and so on. The Waiter brings them orders, and they do the heavy lifting.'),
        ('Microsoft AI Foundry (The Master Brain)', 'Our application relies on Artificial Intelligence. Instead of building a brain from scratch, our application talks securely to Microsoft\'s AI Foundry. Whenever our specialized chefs need to understand complex language, they phone the Master Brain in the cloud for the answer.'),
    ]
    
    for title, desc in parts:
        p_part = doc.add_paragraph(style='List Bullet')
        p_part.add_run(title + ': ').bold = True
        p_part.add_run(desc)

    # Deployment Process
    doc.add_heading('4. Step-by-Step: Moving in (The Deployment Process)', level=1)
    doc.add_paragraph('When we have a new, improved version of our application, here is exactly how we deliver it to Microsoft Azure:')
    
    steps = [
        ('Step 1: Packing the Boxes (Zipping the Code)', 'We take all of the thousands of text files that make up our application and compress them into a single file (a ZIP file). This makes it easy to transport.'),
        ('Step 2: Driving the Moving Truck to Azure', 'Using a secure connection, we upload this single ZIP file directly into our rented "Azure App Service".'),
        ('Step 3: Unpacking and Buying Ingredients (The Build Process)', 'Once Azure receives the box, it unpacks it. Then, a robot worker called "Oryx" reads our recipe list (a file called package.json). Oryx goes out to the internet and automatically downloads all the standard ingredients (tools and libraries) our code needs to run properly. This process is fully automated.'),
        ('Step 4: Hiring the Waiter and Chefs (Startup)', 'Once everything is unpacked and installed, Azure runs our "startup" script. This script boots up the 8 Specialized Chefs one by one, and finally opens the doors by starting the Waiter (API Gateway).'),
        ('Step 5: The Health Inspector', 'Before fully opening to the public, the system checks a special "/health" heartbeat. It asks each of the 8 chefs, "Are you awake?" Only when all 8 say "Yes", the application is marked as healthy and ready to serve users.')
    ]
    
    for title, desc in steps:
        p_step = doc.add_paragraph(style='List Number')
        p_step.add_run(title + '\n').bold = True
        p_step.add_run(desc)
        
    # Configuration and Security
    doc.add_heading('5. Security and Configuration (The Vault)', level=1)
    doc.add_paragraph('Our application needs highly sensitive passwords (called "API Keys") to talk to the AI Master Brain. We NEVER pack these passwords into the ZIP file. Instead, we type these passwords directly into a secure vault inside the Microsoft Azure portal. When the application starts up, it reads from this vault locally, ensuring that no hacker can steal our keys in transit.')

    # Summary
    doc.add_heading('6. Why Do We Use This Process?', level=1)
    p_summary = doc.add_paragraph()
    p_summary.add_run('This process ensures that our application is ')
    p_summary.add_run('Reliable, Secure, and Easy to Update').bold = True
    p_summary.add_run('. If the application crashes, Azure automatically restarts it safely. If we want to release a new feature, we just pack a new ZIP file and upload it. It takes less than 10 minutes to deliver a brand new version of the app to users worldwide.')

    # Save
    doc_path = 'C:\\Users\\pfernandes\\AgentOps\\AgentOps\\Detailed_Simple_Deployment_Guide.docx'
    doc.save(doc_path)
    print(f"Document created successfully at {doc_path}")

if __name__ == '__main__':
    create_doc()
