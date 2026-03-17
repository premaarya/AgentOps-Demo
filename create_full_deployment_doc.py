from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_cell(cell, text, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(10)

def create_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # ---- TITLE PAGE ----
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_heading("Contract AgentOps", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_heading("Deployment Guide: Process, Challenges, and Solutions", 2)
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")
    meta = doc.add_paragraph("Prepared: March 13, 2026")
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta2 = doc.add_paragraph("Azure Subscription: Contoso (ME-MngEnvMCAP179959)")
    meta2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # ---- TABLE OF CONTENTS ----
    doc.add_heading("Table of Contents", 1)
    toc_items = [
        "1. What is Deployment? (The Simple Explanation)",
        "2. What Are We Deploying? (The Application)",
        "3. Where Does It Run? (The Cloud Infrastructure)",
        "4. Step-by-Step Deployment Process",
        "5. Challenge #1: No Room at the Inn (Region Quota Limits)",
        "6. Challenge #2: Missing Ingredients (Build Not Running)",
        "7. Challenge #3: The App Kept Falling Asleep (Process Keep-Alive)",
        "8. Challenge #4: Duplicate Workers Causing Confusion",
        "9. Challenge #5: The App Was Too Slow to Wake Up (Startup Timeout)",
        "10. Challenge #6: Secret Passwords and Security",
        "11. The Final Working Setup",
        "12. How to Update the App in the Future",
        "13. Glossary of Terms",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # ---- SECTION 1 ----
    doc.add_heading("1. What is Deployment? (The Simple Explanation)", 1)
    doc.add_paragraph(
        'Imagine you have written a recipe book on your home computer. '
        'Only you can read it. "Deployment" is the process of printing that book '
        'and placing copies in every bookstore in the world so anyone can buy and read it.'
    )
    doc.add_paragraph(
        'In our case, the "recipe book" is the Contract AgentOps application (a website with AI capabilities), '
        'and the "bookstore" is Microsoft Azure, a massive network of computers owned by Microsoft that anyone '
        'on the internet can access 24 hours a day, 7 days a week.'
    )
    doc.add_paragraph(
        'The goal of deployment is simple: take code that works on a developer\'s laptop and make it '
        'available to real users on the internet, reliably and securely.'
    )

    # ---- SECTION 2 ----
    doc.add_heading("2. What Are We Deploying? (The Application)", 1)
    doc.add_paragraph(
        'Contract AgentOps is an AI-powered legal contract management system. '
        'It can read contracts, extract key information, check for compliance issues, '
        'track changes over time, and provide audit trails. Think of it as a team of virtual '
        'legal assistants working together.'
    )
    doc.add_heading("The Team of Virtual Assistants", 2)
    doc.add_paragraph("The application is made up of several specialized workers:")

    team = [
        ("The Dashboard", "The website you see and click on. It is the \"front door\" of the application."),
        ("The API Gateway (The Receptionist)", "When you click a button, the Gateway receives your request and routes it to the right specialist."),
        ("8 Specialized MCP Servers (The Experts)", "Each one handles a specific job:"),
    ]
    for name, desc in team:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(name + ": ").bold = True
        p.add_run(desc)

    experts = [
        "Contract Intake - Receives and registers new contracts",
        "Contract Extraction - Reads and pulls out key data from contracts",
        "Contract Compliance - Checks contracts against legal rules",
        "Contract Workflow - Manages approval processes",
        "Contract Audit - Keeps a record of every action taken",
        "Contract Evaluation - Scores and grades contracts",
        "Contract Drift - Detects when contracts change unexpectedly",
        "Contract Feedback - Collects and processes user feedback",
    ]
    for expert in experts:
        doc.add_paragraph(expert, style="List Bullet 2")

    p_ai = doc.add_paragraph(style="List Bullet")
    p_ai.add_run("Microsoft AI Foundry (The Brain): ").bold = True
    p_ai.add_run(
        "The AI engine that gives our experts their intelligence. "
        "It is hosted separately by Microsoft and our app calls it over the internet using a secure key."
    )

    # ---- SECTION 3 ----
    doc.add_heading("3. Where Does It Run? (The Cloud Infrastructure)", 1)
    doc.add_paragraph(
        'We use Microsoft Azure, which is like renting a fully managed office building. '
        'Microsoft handles the electricity, internet, physical security, and hardware maintenance. '
        'We just move our application in and it runs.'
    )

    doc.add_heading("Key Azure Resources We Use", 2)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Resource", "What It Is (Simple)", "Our Specific Name"]
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)

    rows_data = [
        ("Resource Group", "A folder that holds all related resources together", "contract-agentops-eastus-rg"),
        ("App Service Plan", "The size of computer we are renting (like choosing a small/medium/large office)", "contract-agentops-eastus-plan (B1 tier)"),
        ("Web App", "The actual cloud computer running our code", "contract-agentops-eastus"),
        ("AI Foundry", "Microsoft's AI brain that our app calls for intelligence", "kpmg-legal-ai-foundry (East US)"),
        ("AI Model", "The specific AI model we use (GPT-4o)", "gpt-4o (version 2024-11-20)"),
    ]
    for r, (a, b, c) in enumerate(rows_data, start=1):
        set_cell(table.rows[r].cells[0], a)
        set_cell(table.rows[r].cells[1], b)
        set_cell(table.rows[r].cells[2], c)

    # ---- SECTION 4 ----
    doc.add_heading("4. Step-by-Step Deployment Process", 1)
    doc.add_paragraph("Here is the exact sequence of events that happens every time we deploy a new version:")

    steps = [
        (
            "Step 1: Pull the Latest Code",
            "We download the newest version of our code from GitHub (a website where code is stored and shared among developers). "
            "This ensures we are deploying the most up-to-date version.",
            'Command: git pull origin main',
        ),
        (
            "Step 2: Apply Fixes",
            "Sometimes the freshly downloaded code needs small adjustments to work correctly in the cloud environment "
            "(more on this in the Challenges sections below). We make these adjustments on our laptop before packaging.",
            None,
        ),
        (
            "Step 3: Pack the Code into a ZIP File",
            "We compress all the application files into a single ZIP archive. This is like packing all your belongings "
            "into moving boxes before a house move. We deliberately exclude unnecessary files (like developer tools "
            "and documentation) to keep the package small (~0.35 MB).",
            'Command: tar -cf deploy.zip [files]',
        ),
        (
            "Step 4: Upload the ZIP to Azure",
            "We securely upload the ZIP file to our Azure Web App. This is like driving the moving truck to the new office.",
            'Command: az webapp deploy --src-path deploy.zip --type zip',
        ),
        (
            "Step 5: Azure Builds the App (Oryx Build)",
            'Once the ZIP arrives, Azure\'s built-in robot called "Oryx" unpacks it, reads our ingredient list '
            '(package.json), and automatically downloads and installs all the libraries and tools the app needs. '
            'This step takes about 2-3 minutes.',
            "Automatic - triggered by SCM_DO_BUILD_DURING_DEPLOYMENT=true",
        ),
        (
            "Step 6: Azure Starts the Application",
            "Azure runs our startup command (npx tsx start.ts). This script boots up the 8 expert servers one by one, "
            "waits for each to report healthy, then starts the Gateway.",
            'Startup command: npx tsx start.ts',
        ),
        (
            "Step 7: Health Check",
            "Azure continuously pings a special /api/v1/health address. The app responds with the status of all 8 "
            "expert servers. Only when all 8 report 'online' does Azure consider the deployment successful and "
            "start sending real user traffic to it.",
            'Health endpoint: https://contract-agentops-eastus.azurewebsites.net/api/v1/health',
        ),
    ]
    for title, desc, cmd in steps:
        p_step = doc.add_paragraph()
        p_step.add_run(title + "\n").bold = True
        p_step.add_run(desc)
        if cmd:
            doc.add_paragraph(cmd, style="List Bullet")
    doc.add_page_break()

    # ---- CHALLENGES ----
    doc.add_heading("5. Challenge #1: No Room at the Inn (Region Quota Limits)", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "Our AI Foundry brain lives in Microsoft's East US data center. Ideally, we wanted our application "
        "to live in the same building (same region) so they could talk to each other as fast as possible. "
        "However, Microsoft told us: 'Sorry, East US is completely full. You cannot rent any computer space here.' "
        "This is called a quota limit -- like a hotel with no vacancies."
    )
    doc.add_heading("What We Tried", 2)
    doc.add_paragraph("We checked 7 different US regions to see which ones had availability:")
    avail = [
        ("East US", "FULL (0 quota)"),
        ("South Central US", "FULL"),
        ("West US", "FULL"),
        ("East US 2", "AVAILABLE"),
        ("Central US", "AVAILABLE"),
        ("West US 2", "AVAILABLE"),
        ("West US 3", "AVAILABLE"),
    ]
    t2 = doc.add_table(rows=len(avail) + 1, cols=2)
    t2.style = "Light Grid Accent 1"
    set_cell(t2.rows[0].cells[0], "Region", bold=True)
    set_cell(t2.rows[0].cells[1], "Status", bold=True)
    for i, (region, status) in enumerate(avail, 1):
        set_cell(t2.rows[i].cells[0], region)
        set_cell(t2.rows[i].cells[1], status)

    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        "We chose East US 2, which is the closest available neighbor to East US. "
        "The AI brain is in East US and our app is in East US 2 -- they are in nearby data centers, "
        "so communication is still very fast (just a few milliseconds of extra travel time)."
    )

    # ---- CHALLENGE 2 ----
    doc.add_heading("6. Challenge #2: Missing Ingredients (Build Not Running)", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "When we first uploaded the ZIP file, the application crashed immediately. "
        "The reason: Azure received our code but never installed the ingredients (libraries) it needed. "
        "It was like shipping a cake recipe to a kitchen that had no flour, sugar, or eggs."
    )
    doc.add_paragraph(
        "Technically, the upload method we used (called 'OneDeploy') simply extracts the ZIP and runs the app. "
        "It does NOT automatically run the installation step (npm install) unless you explicitly tell it to."
    )
    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        "We added a special setting to the Azure configuration:"
    )
    p_fix = doc.add_paragraph(style="List Bullet")
    p_fix.add_run("SCM_DO_BUILD_DURING_DEPLOYMENT = true").bold = True
    doc.add_paragraph(
        "This single setting tells Azure: 'After you receive the ZIP, please also run the installation step.' "
        "With this enabled, Azure's Oryx robot automatically reads our recipe list and downloads all necessary "
        "ingredients. The build step now takes about 2-3 minutes but the app starts correctly."
    )

    # ---- CHALLENGE 3 ----
    doc.add_heading("7. Challenge #3: The App Kept Falling Asleep (Process Keep-Alive)", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "After fixing the build, the app would start successfully -- we could see the 'Ready!' message in the "
        "logs -- but then it would shut itself down after about 60 seconds. Azure would then mark the site as "
        "'failed' because it disappeared."
    )
    doc.add_paragraph(
        "The root cause: In the programming language we use (Node.js/TypeScript), a program automatically exits "
        "when it has nothing left to do. Our startup script would start all 9 workers, print 'Ready!', and "
        "then... the main script had nothing else on its to-do list, so Node.js assumed the job was done and "
        "shut everything down. It is like a project manager who finishes assigning tasks, then leaves the building "
        "and accidentally locks everyone else out."
    )
    doc.add_heading("What We Tried First (Did Not Work)", 2)
    doc.add_paragraph(
        'We initially tried: await new Promise(() => {}). This was supposed to create an "infinite wait" that '
        "would keep the program alive. However, Node.js is smart enough to realize that this promise will "
        "never resolve AND there is nothing else happening, so it exits anyway."
    )
    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        'We replaced it with: setInterval(() => {}, 60000). This creates a recurring timer that fires every '
        "60 seconds. Because Node.js always has this upcoming timer on its schedule, it never considers itself "
        '"done" and stays alive indefinitely. It is like giving the project manager a recurring meeting -- '
        "they will never leave the building because there is always another meeting coming up."
    )

    # ---- CHALLENGE 4 ----
    doc.add_heading("8. Challenge #4: Duplicate Workers Causing Confusion", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "The application has a function called 'waitForHealth' that checks if each expert server is awake. "
        "After pulling the latest code, we discovered that this function existed in TWO places:"
    )
    places = [
        "Imported from a separate file (startup/health.ts) -- the official, well-tested version",
        "Also defined directly inside the main startup script (start.ts) -- an older, outdated copy",
    ]
    for place in places:
        doc.add_paragraph(place, style="List Bullet")
    doc.add_paragraph(
        "When two workers have the same name but do the job differently, chaos ensues. The program would crash "
        "because it could not decide which version to use, and the two versions expected different inputs."
    )
    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        "We simply deleted the old duplicate copy from start.ts, leaving only the official imported version. "
        "One worker, one job, no confusion."
    )

    # ---- CHALLENGE 5 ----
    doc.add_heading("9. Challenge #5: The App Was Too Slow to Wake Up (Startup Timeout)", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "Azure gives every application a limited amount of time to start up and respond to health checks. "
        "The default is about 230 seconds (~4 minutes). Our application needs to boot 9 separate programs "
        "in sequence, and on a basic (B1) tier computer, this can take longer than 4 minutes. "
        "Azure would give up waiting and declare the app 'dead' even though it was still starting."
    )
    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        "We increased the patience setting:"
    )
    p_timeout = doc.add_paragraph(style="List Bullet")
    p_timeout.add_run("WEBSITES_CONTAINER_START_TIME_LIMIT = 600").bold = True
    doc.add_paragraph(
        "This tells Azure: 'Please wait up to 10 minutes before giving up.' "
        "This gave the application plenty of time to start all 9 workers at its own pace."
    )

    # ---- CHALLENGE 6 ----
    doc.add_heading("10. Challenge #6: Secret Passwords and Security", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "Our application needs a secret password (API Key) to talk to the AI Foundry brain. "
        "If we packed this password into the ZIP file, anyone who intercepted the file could steal it "
        "and run up charges on our AI account, or worse, access sensitive data."
    )
    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        "We NEVER put passwords in the code. Instead, we store them as 'App Settings' directly inside "
        "Azure's secure configuration panel. When the application starts, it reads these settings from "
        "the local environment (like checking a locked safe in the office) rather than carrying them "
        "in the moving boxes."
    )
    doc.add_paragraph("The settings we configured:")
    settings = [
        ("FOUNDRY_ENDPOINT", "The address of the AI brain"),
        ("FOUNDRY_API_KEY", "The secret password to talk to the AI brain"),
        ("FOUNDRY_MODEL", "Which AI model to use (gpt-4o)"),
        ("GATEWAY_PORT", "Which door number the app listens on (8080)"),
        ("DEMO_MODE", "Whether to use real AI or simulated responses (live)"),
        ("SCM_DO_BUILD_DURING_DEPLOYMENT", "Tells Azure to install ingredients (true)"),
        ("WEBSITES_CONTAINER_START_TIME_LIMIT", "How long Azure waits for startup (600 seconds)"),
    ]
    t3 = doc.add_table(rows=len(settings) + 1, cols=2)
    t3.style = "Light Grid Accent 1"
    set_cell(t3.rows[0].cells[0], "Setting Name", bold=True)
    set_cell(t3.rows[0].cells[1], "What It Means", bold=True)
    for i, (name, meaning) in enumerate(settings, 1):
        set_cell(t3.rows[i].cells[0], name)
        set_cell(t3.rows[i].cells[1], meaning)
    doc.add_page_break()

    # ---- FINAL SETUP ----
    doc.add_heading("11. The Final Working Setup", 1)
    doc.add_paragraph("After solving all of the above challenges, here is the successfully deployed application:")

    final = [
        ("Application URL", "https://contract-agentops-eastus.azurewebsites.net/"),
        ("Health Check URL", "https://contract-agentops-eastus.azurewebsites.net/api/v1/health"),
        ("Region", "East US 2"),
        ("Computer Size", "B1 (Basic tier, Linux)"),
        ("Runtime", "Node.js 20 LTS"),
        ("AI Model", "GPT-4o via Azure AI Foundry in East US"),
        ("Status", "All 8 expert servers ONLINE, Gateway HEALTHY"),
    ]
    t4 = doc.add_table(rows=len(final) + 1, cols=2)
    t4.style = "Light Grid Accent 1"
    set_cell(t4.rows[0].cells[0], "Item", bold=True)
    set_cell(t4.rows[0].cells[1], "Value", bold=True)
    for i, (k, v) in enumerate(final, 1):
        set_cell(t4.rows[i].cells[0], k)
        set_cell(t4.rows[i].cells[1], v)

    doc.add_paragraph("")
    doc.add_paragraph("The health check returns:")
    doc.add_paragraph(
        '{\n'
        '  "status": "ok",\n'
        '  "mode": "live",\n'
        '  "servers": {\n'
        '    "contract-intake-mcp": "online",\n'
        '    "contract-extraction-mcp": "online",\n'
        '    "contract-compliance-mcp": "online",\n'
        '    "contract-workflow-mcp": "online",\n'
        '    "contract-audit-mcp": "online",\n'
        '    "contract-eval-mcp": "online",\n'
        '    "contract-drift-mcp": "online",\n'
        '    "contract-feedback-mcp": "online"\n'
        '  }\n'
        '}'
    )

    # ---- HOW TO UPDATE ----
    doc.add_heading("12. How to Update the App in the Future", 1)
    doc.add_paragraph("Now that the initial setup is done, future updates are much simpler:")
    future_steps = [
        "Pull the latest code from GitHub (git pull origin main)",
        "Re-apply the known fixes if the upstream code has reverted them (build scripts, keep-alive, tsx in dependencies)",
        "Pack a new ZIP (excluding node_modules and .git)",
        "Upload to Azure (az webapp deploy --src-path deploy.zip --type zip)",
        "Wait 3-5 minutes for build + startup",
        "Verify health at /api/v1/health",
    ]
    for s in future_steps:
        doc.add_paragraph(s, style="List Number")

    # ---- GLOSSARY ----
    doc.add_heading("13. Glossary of Terms", 1)
    glossary = [
        ("Azure", "Microsoft's cloud computing platform -- a network of data centers around the world that you can rent computer space from."),
        ("App Service", "An Azure product that runs web applications without you managing the underlying computer hardware."),
        ("Resource Group", "A container (like a folder) in Azure that groups related resources together for easy management."),
        ("Oryx", "Azure's automatic build system that reads your project files and installs all required software libraries."),
        ("npm install", "The command that downloads and installs all the code libraries your application depends on."),
        ("ZIP Deploy", "A method of deploying by uploading a compressed archive of your code to Azure."),
        ("API Key", "A secret password string used to authenticate with an external service (like AI Foundry)."),
        ("Health Check", "An automated test that asks 'Is the application alive and working?' at regular intervals."),
        ("Quota", "A limit on how much of a cloud resource you are allowed to use in a particular region."),
        ("Region", "A physical location of a Microsoft data center (e.g., East US, West Europe)."),
        ("MCP Server", "Model Context Protocol server -- a specialized mini-application that handles one specific type of work."),
        ("Gateway", "The single entry point that receives all user requests and routes them to the correct internal service."),
        ("Node.js", "The programming language runtime our application uses (JavaScript/TypeScript)."),
        ("tsx", "A tool that lets us run TypeScript code directly without a separate compilation step."),
        ("Git / GitHub", "A version control system and website for storing, sharing, and collaborating on code."),
    ]
    t5 = doc.add_table(rows=len(glossary) + 1, cols=2)
    t5.style = "Light Grid Accent 1"
    set_cell(t5.rows[0].cells[0], "Term", bold=True)
    set_cell(t5.rows[0].cells[1], "Definition", bold=True)
    for i, (term, defn) in enumerate(glossary, 1):
        set_cell(t5.rows[i].cells[0], term)
        set_cell(t5.rows[i].cells[1], defn)

    # Save
    doc_path = "C:\\Users\\pfernandes\\AgentOps\\AgentOps\\Contract_AgentOps_Deployment_Guide_Full.docx"
    doc.save(doc_path)
    print(f"Document created successfully at {doc_path}")

if __name__ == "__main__":
    create_doc()
