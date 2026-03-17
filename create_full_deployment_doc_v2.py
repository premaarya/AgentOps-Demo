from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_cell(cell, text, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(10)

def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    return p

def create_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # ---- TITLE PAGE ----
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_heading("Contract AgentOps", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_heading("Deployment Guide: Process, Challenges, and Solutions", 2)
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")
    meta = doc.add_paragraph("Prepared: March 13, 2026")
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta2 = doc.add_paragraph("Azure Subscription: Contoso (ME-MngEnvMCAP179959)")
    meta2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # ---- TABLE OF CONTENTS ----
    doc.add_heading("Table of Contents", 1)
    toc_items = [
        "1. What is Deployment? (The Simple Explanation)",
        "2. What Are We Deploying? (The Application)",
        "3. Where Does It Run? (The Cloud Infrastructure)",
        "4. Setting Up the Cloud: Creating Azure Resources Step-by-Step",
        "5. Configuring the Web App (App Settings and Startup)",
        "6. The Deployment Process: Uploading and Running the Code",
        "7. Challenge #1: No Room at the Inn (Region Quota Limits)",
        "8. Challenge #2: Missing Ingredients (Build Not Running)",
        "9. Challenge #3: The App Kept Falling Asleep (Process Keep-Alive)",
        "10. Challenge #4: Duplicate Workers Causing Confusion",
        "11. Challenge #5: The App Was Too Slow to Wake Up (Startup Timeout)",
        "12. Challenge #6: Secret Passwords and Security",
        "13. The Final Working Setup",
        "14. How to Update the App in the Future",
        "15. Glossary of Terms",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # ---- SECTION 1 ----
    doc.add_heading("1. What is Deployment? (The Simple Explanation)", 1)
    doc.add_paragraph(
        'Imagine you have written a recipe book on your home computer. '
        'Only you can read it. "Deployment" is the process of printing that book '
        'and placing copies in every bookstore in the world so anyone can buy and read it.'
    )
    doc.add_paragraph(
        'In our case, the "recipe book" is the Contract AgentOps application (a website with AI capabilities), '
        'and the "bookstore" is Microsoft Azure, a massive network of computers owned by Microsoft that anyone '
        'on the internet can access 24 hours a day, 7 days a week.'
    )
    doc.add_paragraph(
        'The goal of deployment is simple: take code that works on a developer\'s laptop and make it '
        'available to real users on the internet, reliably and securely.'
    )

    # ---- SECTION 2 ----
    doc.add_heading("2. What Are We Deploying? (The Application)", 1)
    doc.add_paragraph(
        'Contract AgentOps is an AI-powered legal contract management system. '
        'It can read contracts, extract key information, check for compliance issues, '
        'track changes over time, and provide audit trails. Think of it as a team of virtual '
        'legal assistants working together.'
    )
    doc.add_heading("The Team of Virtual Assistants", 2)
    doc.add_paragraph("The application is made up of several specialized workers:")

    team = [
        ("The Dashboard", "The website you see and click on. It is the \"front door\" of the application."),
        ("The API Gateway (The Receptionist)", "When you click a button, the Gateway receives your request and routes it to the right specialist."),
        ("8 Specialized MCP Servers (The Experts)", "Each one handles a specific job:"),
    ]
    for name, desc in team:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(name + ": ").bold = True
        p.add_run(desc)

    experts = [
        "Contract Intake - Receives and registers new contracts",
        "Contract Extraction - Reads and pulls out key data from contracts",
        "Contract Compliance - Checks contracts against legal rules",
        "Contract Workflow - Manages approval processes",
        "Contract Audit - Keeps a record of every action taken",
        "Contract Evaluation - Scores and grades contracts",
        "Contract Drift - Detects when contracts change unexpectedly",
        "Contract Feedback - Collects and processes user feedback",
    ]
    for expert in experts:
        doc.add_paragraph(expert, style="List Bullet 2")

    p_ai = doc.add_paragraph(style="List Bullet")
    p_ai.add_run("Microsoft AI Foundry (The Brain): ").bold = True
    p_ai.add_run(
        "The AI engine that gives our experts their intelligence. "
        "It is hosted separately by Microsoft and our app calls it over the internet using a secure key."
    )

    # ---- SECTION 3 ----
    doc.add_heading("3. Where Does It Run? (The Cloud Infrastructure)", 1)
    doc.add_paragraph(
        'We use Microsoft Azure, which is like renting a fully managed office building. '
        'Microsoft handles the electricity, internet, physical security, and hardware maintenance. '
        'We just move our application in and it runs.'
    )

    doc.add_heading("Key Azure Resources We Use", 2)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Resource", "What It Is (Simple)", "Our Specific Name"]
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)

    rows_data = [
        ("Resource Group", "A folder that holds all related resources together", "contract-agentops-eastus-rg"),
        ("App Service Plan", "The size of computer we are renting (like choosing a small/medium/large office)", "contract-agentops-eastus-plan (B1 tier)"),
        ("Web App", "The actual cloud computer running our code", "contract-agentops-eastus"),
        ("AI Foundry", "Microsoft's AI brain that our app calls for intelligence", "kpmg-legal-ai-foundry (East US)"),
        ("AI Model", "The specific AI model we use (GPT-4o)", "gpt-4o (version 2024-11-20)"),
    ]
    for r, (a, b, c) in enumerate(rows_data, start=1):
        set_cell(table.rows[r].cells[0], a)
        set_cell(table.rows[r].cells[1], b)
        set_cell(table.rows[r].cells[2], c)

    doc.add_paragraph("")
    doc.add_paragraph(
        "The next section explains exactly how each of these resources was created, "
        "step by step, using the Azure command-line tool."
    )
    doc.add_page_break()

    # ---- SECTION 4: CREATING AZURE RESOURCES (NEW DETAILED SECTION) ----
    doc.add_heading("4. Setting Up the Cloud: Creating Azure Resources Step-by-Step", 1)
    doc.add_paragraph(
        "Before we can upload any code, we first need to build the \"office\" that our application will live in. "
        "This section walks through every step of creating the Azure infrastructure from scratch. "
        "All of this is done using the Azure CLI (Command Line Interface) -- a tool that lets us type commands "
        "on our laptop to create and manage cloud resources remotely."
    )

    # 4.1 Prerequisites
    doc.add_heading("4.1 Prerequisites: What You Need Before Starting", 2)
    doc.add_paragraph("Before creating any resources, you need the following tools installed on your laptop:")
    prereqs = [
        ("Azure CLI (az)", "The command-line tool for managing Azure. Install from https://aka.ms/installazurecliwindows"),
        ("A Microsoft Azure account", "You need a valid subscription with permissions to create resources."),
        ("PowerShell or Terminal", "A window where you type commands."),
    ]
    for name, desc in prereqs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(name + ": ").bold = True
        p.add_run(desc)

    # 4.2 Login
    doc.add_heading("4.2 Step 1: Logging In to Azure", 2)
    doc.add_paragraph(
        "The first thing we do is prove to Microsoft that we are who we say we are. "
        "This is called authentication. We run the login command, which opens a web browser "
        "where you enter your Microsoft account credentials."
    )
    doc.add_heading("The Command:", 3)
    add_code_block(doc, "az login --tenant 5fb8f011-cf24-4891-a2a9-7953d7e290bd")
    doc.add_paragraph(
        "The --tenant part tells Azure which organization we belong to (in our case, the Contoso tenant). "
        "After logging in, Azure remembers our identity for future commands."
    )
    doc.add_heading("What is a Tenant?", 3)
    doc.add_paragraph(
        "A tenant is like a company's membership card for Azure. Every organization has one unique tenant ID. "
        "It ensures that only people from your organization can access your company's resources."
    )

    # 4.2b Set Subscription
    doc.add_heading("4.3 Step 2: Selecting the Right Subscription", 2)
    doc.add_paragraph(
        "A single organization can have multiple Azure subscriptions (think of them as separate billing accounts). "
        "We need to tell Azure which one to charge for the resources we are about to create."
    )
    doc.add_heading("The Command:", 3)
    add_code_block(doc, "az account set --subscription b83ea59a-7cfb-4709-8b96-2fa54b0af84f")
    doc.add_paragraph(
        "This selects our specific subscription: ME-MngEnvMCAP179959-clgerola-1. "
        "All resources we create from this point forward will be billed to this subscription."
    )

    # 4.3 Resource Group
    doc.add_heading("4.4 Step 3: Creating the Resource Group (The Folder)", 2)
    doc.add_paragraph(
        "Before creating individual resources, we need a container to hold them all together. "
        "In Azure, this container is called a Resource Group. Think of it as an empty folder on your desktop "
        "where you will save all the files for one project."
    )
    doc.add_heading("The Command:", 3)
    add_code_block(doc, "az group create --name contract-agentops-eastus-rg --location eastus2")
    doc.add_heading("What Each Part Means:", 3)
    parts = [
        ("az group create", "This tells Azure: 'Please create a new Resource Group.'"),
        ("--name contract-agentops-eastus-rg", "The name we chose for our folder. The '-rg' suffix is a naming convention that stands for 'resource group.'"),
        ("--location eastus2", "The physical location of the data center. We chose East US 2 because East US was full (see Challenge #1)."),
    ]
    for cmd_part, explanation in parts:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(cmd_part).bold = True
        p.add_run(" -- " + explanation)

    doc.add_paragraph(
        "After running this command, Azure creates an empty folder named 'contract-agentops-eastus-rg' in the "
        "East US 2 data center. It takes just a few seconds."
    )

    # 4.4 App Service Plan
    doc.add_heading("4.5 Step 4: Creating the App Service Plan (Choosing the Computer Size)", 2)
    doc.add_paragraph(
        "Now we need to decide what size of computer to rent. In Azure, this choice is called an App Service Plan. "
        "It determines how much processing power (CPU), memory (RAM), and storage our application gets."
    )
    doc.add_paragraph("Think of it like choosing a rental car:")
    car_analogy = [
        ("Free/Shared tier (F1)", "A bicycle. Free, but very slow and limited."),
        ("Basic tier (B1) -- What We Use", "A small sedan. Good enough for testing and demos. About $13/month."),
        ("Standard tier (S1)", "An SUV. Good for real production use with moderate traffic."),
        ("Premium tier (P1v3)", "A sports car. Fast, powerful, and more expensive."),
    ]
    for tier, analogy in car_analogy:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(tier + ": ").bold = True
        p.add_run(analogy)

    doc.add_heading("The Command:", 3)
    add_code_block(doc,
        'az appservice plan create \\\n'
        '    --name contract-agentops-eastus-plan \\\n'
        '    --resource-group contract-agentops-eastus-rg \\\n'
        '    --sku B1 \\\n'
        '    --is-linux'
    )
    doc.add_heading("What Each Part Means:", 3)
    plan_parts = [
        ("az appservice plan create", "Tells Azure: 'Create a new computer rental plan.'"),
        ("--name contract-agentops-eastus-plan", "The name we give to this plan."),
        ("--resource-group contract-agentops-eastus-rg", "Put it inside the folder we created in Step 3."),
        ("--sku B1", "The size of the computer we want. B1 = Basic tier with 1 CPU core and 1.75 GB of RAM."),
        ("--is-linux", "We want a Linux computer (not Windows). Linux is the industry standard for running Node.js web applications and is slightly cheaper."),
    ]
    for cmd_part, explanation in plan_parts:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(cmd_part).bold = True
        p.add_run(" -- " + explanation)

    doc.add_paragraph(
        "After this command, Azure reserves a small Linux computer for us in the East US 2 data center. "
        "The computer is now waiting for an application to be installed on it."
    )

    # 4.5 Web App
    doc.add_heading("4.6 Step 5: Creating the Web App (Installing the Office Space)", 2)
    doc.add_paragraph(
        "The App Service Plan is the building, and now we need to create the actual office space (the Web App) inside it. "
        "The Web App is the specific entity that will hold our code, have its own URL, and serve users."
    )
    doc.add_heading("The Command:", 3)
    add_code_block(doc,
        'az webapp create \\\n'
        '    --name contract-agentops-eastus \\\n'
        '    --resource-group contract-agentops-eastus-rg \\\n'
        '    --plan contract-agentops-eastus-plan \\\n'
        '    --runtime "NODE|20-lts"'
    )
    doc.add_heading("What Each Part Means:", 3)
    webapp_parts = [
        ("az webapp create", "Tells Azure: 'Create a new Web Application.'"),
        ("--name contract-agentops-eastus", "The name of our web app. This name becomes part of the public URL: https://contract-agentops-eastus.azurewebsites.net/"),
        ("--resource-group contract-agentops-eastus-rg", "Put it in our project folder."),
        ("--plan contract-agentops-eastus-plan", "Run it on the B1 computer we just reserved."),
        ('--runtime "NODE|20-lts"', "Tells Azure what programming language our app uses: Node.js version 20 (Long Term Support). This ensures the right software is pre-installed on the computer."),
    ]
    for cmd_part, explanation in webapp_parts:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(cmd_part).bold = True
        p.add_run(" -- " + explanation)

    doc.add_heading("A Small Hiccup We Encountered:", 3)
    doc.add_paragraph(
        'When running this command on Windows using PowerShell, the pipe character (|) inside "NODE|20-lts" '
        "caused problems. PowerShell thought the | was a command separator, not part of the value. "
        "We solved this by using a special PowerShell trick called the stop-parsing token:"
    )
    add_code_block(doc, 'az --% webapp create --name contract-agentops-eastus --resource-group contract-agentops-eastus-rg --plan contract-agentops-eastus-plan --runtime "NODE|20-lts"')
    doc.add_paragraph(
        'The --% tells PowerShell: "Stop trying to interpret what comes next. Just pass it as-is to the az tool." '
        "This is a common issue when using Azure CLI on Windows."
    )

    doc.add_heading("What We Got:", 3)
    doc.add_paragraph("After this command succeeds, Azure gives us:")
    results = [
        "A live web app at https://contract-agentops-eastus.azurewebsites.net/",
        "A default placeholder page (saying 'Your web app is running')",
        "A Kudu management portal at https://contract-agentops-eastus.scm.azurewebsites.net/",
        "HTTPS enabled by default (secure connection, the little lock icon in the browser)",
    ]
    for r in results:
        doc.add_paragraph(r, style="List Bullet")
    doc.add_page_break()

    # ---- SECTION 5: CONFIGURING THE WEB APP ----
    doc.add_heading("5. Configuring the Web App (App Settings and Startup)", 1)
    doc.add_paragraph(
        "Now that the office space exists, we need to furnish it before moving in. "
        "This means telling the app where to find the AI brain, setting passwords, "
        "and telling Azure how to start our application."
    )

    # 5.1 App Settings
    doc.add_heading("5.1 Setting Environment Variables (The App's Configuration File)", 2)
    doc.add_paragraph(
        "Environment variables are like a set of instructions pinned to the office wall. "
        "When the application starts, it reads these instructions to know how to behave. "
        "We set them using the Azure CLI."
    )
    doc.add_heading("The Command:", 3)
    add_code_block(doc,
        'az webapp config appsettings set \\\n'
        '    --name contract-agentops-eastus \\\n'
        '    --resource-group contract-agentops-eastus-rg \\\n'
        '    --settings \\\n'
        '        FOUNDRY_ENDPOINT=https://kpmg-legal-ai-foundry.openai.azure.com/ \\\n'
        '        FOUNDRY_API_KEY=<secret-key-hidden> \\\n'
        '        FOUNDRY_MODEL=gpt-4o \\\n'
        '        GATEWAY_PORT=8080 \\\n'
        '        DEMO_MODE=live \\\n'
        '        MCP_BASE_PORT=9001 \\\n'
        '        LOG_LEVEL=INFO \\\n'
        '        SCM_DO_BUILD_DURING_DEPLOYMENT=true \\\n'
        '        WEBSITES_CONTAINER_START_TIME_LIMIT=600'
    )
    doc.add_heading("What Each Setting Means:", 3)
    settings_detail = [
        ("FOUNDRY_ENDPOINT", "The internet address of the AI brain. Like a phone number our app dials to ask AI questions."),
        ("FOUNDRY_API_KEY", "The secret password that proves our app is allowed to use the AI brain. NEVER shared in code."),
        ("FOUNDRY_MODEL", "Which AI model to use. We chose gpt-4o, one of the most capable models available."),
        ("GATEWAY_PORT=8080", "The 'door number' our app listens on. Azure's load balancer routes incoming traffic to this door. Port 8080 is the standard for Azure App Service."),
        ("DEMO_MODE=live", "Controls whether the app uses real AI or fake/simulated responses. 'live' means real AI."),
        ("MCP_BASE_PORT=9001", "The first door number for our 8 expert servers. They use ports 9001 through 9008."),
        ("LOG_LEVEL=INFO", "How much detail the app writes to its diary (logs). INFO is a good balance."),
        ("SCM_DO_BUILD_DURING_DEPLOYMENT=true", "Tells Azure to install all code libraries after receiving our ZIP file (see Challenge #2)."),
        ("WEBSITES_CONTAINER_START_TIME_LIMIT=600", "Tells Azure to wait up to 10 minutes for the app to start (see Challenge #5)."),
    ]
    t_settings = doc.add_table(rows=len(settings_detail) + 1, cols=2)
    t_settings.style = "Light Grid Accent 1"
    set_cell(t_settings.rows[0].cells[0], "Setting", bold=True)
    set_cell(t_settings.rows[0].cells[1], "What It Means (Plain English)", bold=True)
    for i, (name, meaning) in enumerate(settings_detail, 1):
        set_cell(t_settings.rows[i].cells[0], name)
        set_cell(t_settings.rows[i].cells[1], meaning)

    # 5.2 Startup Command
    doc.add_heading("5.2 Setting the Startup Command", 2)
    doc.add_paragraph(
        "By default, Azure tries to start Node.js apps by running a file called 'server.js.' "
        "Our app uses a different file (start.ts) and a different tool (tsx) to run it. "
        "So we need to tell Azure exactly how to start our application."
    )
    doc.add_heading("The Command:", 3)
    add_code_block(doc,
        'az webapp config set \\\n'
        '    --name contract-agentops-eastus \\\n'
        '    --resource-group contract-agentops-eastus-rg \\\n'
        '    --startup-file "npx tsx start.ts"'
    )
    doc.add_paragraph("This tells Azure: 'When starting the app, run the command npx tsx start.ts.' Breaking it down:")
    startup_parts = [
        ("npx", "A tool that runs programs installed in the project."),
        ("tsx", "A TypeScript executor -- it reads our TypeScript code and runs it directly without needing a separate compilation step."),
        ("start.ts", "Our main startup script that boots all 8 expert servers and the Gateway."),
    ]
    for part, desc in startup_parts:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(part + ": ").bold = True
        p.add_run(desc)
    doc.add_page_break()

    # ---- SECTION 6: DEPLOYMENT PROCESS ----
    doc.add_heading("6. The Deployment Process: Uploading and Running the Code", 1)
    doc.add_paragraph("With the infrastructure built and configured, here is how we actually deliver the code:")

    steps = [
        (
            "Step 1: Pull the Latest Code",
            "We download the newest version of our code from GitHub (a website where code is stored and shared among developers). "
            "This ensures we are deploying the most up-to-date version.",
            'git pull origin main',
        ),
        (
            "Step 2: Apply Known Fixes",
            "Sometimes the freshly downloaded code needs small adjustments to work correctly in the cloud environment "
            "(more on this in the Challenges sections below). We make these adjustments on our laptop before packaging.",
            None,
        ),
        (
            "Step 3: Pack the Code into a ZIP File",
            "We compress all the application files into a single ZIP archive. This is like packing all your belongings "
            "into moving boxes before a house move. We deliberately exclude unnecessary files (like developer tools, "
            "node_modules, .git history, and documentation) to keep the package small (~0.35 MB).",
            'tar -cf deploy.zip --exclude=node_modules --exclude=.git -C AgentOps -a .',
        ),
        (
            "Step 4: Upload the ZIP to Azure",
            "We securely upload the ZIP file to our Azure Web App. This is like driving the moving truck to the new office.",
            'az webapp deploy --name contract-agentops-eastus --resource-group contract-agentops-eastus-rg --src-path deploy.zip --type zip',
        ),
        (
            "Step 5: Azure Builds the App (Oryx Build)",
            'Once the ZIP arrives, Azure\'s built-in robot called "Oryx" unpacks it, reads our ingredient list '
            '(package.json), and automatically downloads and installs all the libraries and tools the app needs. '
            'This step takes about 2-3 minutes. This only happens because we set SCM_DO_BUILD_DURING_DEPLOYMENT=true.',
            None,
        ),
        (
            "Step 6: Azure Starts the Application",
            "Azure runs our startup command (npx tsx start.ts). This script boots up the 8 expert servers one by one, "
            "waits for each to report healthy, then starts the Gateway. This takes about 2-3 minutes on a B1 machine.",
            None,
        ),
        (
            "Step 7: Health Check Verification",
            "Azure continuously pings a special /api/v1/health address. The app responds with the status of all 8 "
            "expert servers. Only when all 8 report 'online' does Azure consider the deployment successful and "
            "start sending real user traffic to it.",
            'https://contract-agentops-eastus.azurewebsites.net/api/v1/health',
        ),
    ]
    for title, desc, cmd in steps:
        p_step = doc.add_paragraph()
        p_step.add_run(title + "\n").bold = True
        p_step.add_run(desc)
        if cmd:
            add_code_block(doc, cmd)
    doc.add_page_break()

    # ---- CHALLENGES ----
    doc.add_heading("7. Challenge #1: No Room at the Inn (Region Quota Limits)", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "Our AI Foundry brain lives in Microsoft's East US data center. Ideally, we wanted our application "
        "to live in the same building (same region) so they could talk to each other as fast as possible. "
        "However, Microsoft told us: 'Sorry, East US is completely full. You cannot rent any computer space here.' "
        "This is called a quota limit -- like a hotel with no vacancies."
    )
    doc.add_heading("What We Tried", 2)
    doc.add_paragraph("We checked 7 different US regions to see which ones had availability:")
    avail = [
        ("East US", "FULL (0 quota)"),
        ("South Central US", "FULL"),
        ("West US", "FULL"),
        ("East US 2", "AVAILABLE"),
        ("Central US", "AVAILABLE"),
        ("West US 2", "AVAILABLE"),
        ("West US 3", "AVAILABLE"),
    ]
    t2 = doc.add_table(rows=len(avail) + 1, cols=2)
    t2.style = "Light Grid Accent 1"
    set_cell(t2.rows[0].cells[0], "Region", bold=True)
    set_cell(t2.rows[0].cells[1], "Status", bold=True)
    for i, (region, status) in enumerate(avail, 1):
        set_cell(t2.rows[i].cells[0], region)
        set_cell(t2.rows[i].cells[1], status)

    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        "We chose East US 2, which is the closest available neighbor to East US. "
        "The AI brain is in East US and our app is in East US 2 -- they are in nearby data centers, "
        "so communication is still very fast (just a few milliseconds of extra travel time)."
    )

    # ---- CHALLENGE 2 ----
    doc.add_heading("8. Challenge #2: Missing Ingredients (Build Not Running)", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "When we first uploaded the ZIP file, the application crashed immediately. "
        "The reason: Azure received our code but never installed the ingredients (libraries) it needed. "
        "It was like shipping a cake recipe to a kitchen that had no flour, sugar, or eggs."
    )
    doc.add_paragraph(
        "Technically, the upload method we used (called 'OneDeploy') simply extracts the ZIP and runs the app. "
        "It does NOT automatically run the installation step (npm install) unless you explicitly tell it to."
    )
    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        "We added a special setting to the Azure configuration:"
    )
    add_code_block(doc, "SCM_DO_BUILD_DURING_DEPLOYMENT = true")
    doc.add_paragraph(
        "This single setting tells Azure: 'After you receive the ZIP, please also run the installation step.' "
        "With this enabled, Azure's Oryx robot automatically reads our recipe list and downloads all necessary "
        "ingredients. The build step now takes about 2-3 minutes but the app starts correctly."
    )

    # ---- CHALLENGE 3 ----
    doc.add_heading("9. Challenge #3: The App Kept Falling Asleep (Process Keep-Alive)", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "After fixing the build, the app would start successfully -- we could see the 'Ready!' message in the "
        "logs -- but then it would shut itself down after about 60 seconds. Azure would then mark the site as "
        "'failed' because it disappeared."
    )
    doc.add_paragraph(
        "The root cause: In the programming language we use (Node.js/TypeScript), a program automatically exits "
        "when it has nothing left to do. Our startup script would start all 9 workers, print 'Ready!', and "
        "then... the main script had nothing else on its to-do list, so Node.js assumed the job was done and "
        "shut everything down. It is like a project manager who finishes assigning tasks, then leaves the building "
        "and accidentally locks everyone else out."
    )
    doc.add_heading("What We Tried First (Did Not Work)", 2)
    doc.add_paragraph(
        'We initially tried: await new Promise(() => {}). This was supposed to create an "infinite wait" that '
        "would keep the program alive. However, Node.js is smart enough to realize that this promise will "
        "never resolve AND there is nothing else happening, so it exits anyway."
    )
    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        'We replaced it with: setInterval(() => {}, 60000). This creates a recurring timer that fires every '
        "60 seconds. Because Node.js always has this upcoming timer on its schedule, it never considers itself "
        '"done" and stays alive indefinitely. It is like giving the project manager a recurring meeting -- '
        "they will never leave the building because there is always another meeting coming up."
    )

    # ---- CHALLENGE 4 ----
    doc.add_heading("10. Challenge #4: Duplicate Workers Causing Confusion", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "The application has a function called 'waitForHealth' that checks if each expert server is awake. "
        "After pulling the latest code, we discovered that this function existed in TWO places:"
    )
    places = [
        "Imported from a separate file (startup/health.ts) -- the official, well-tested version",
        "Also defined directly inside the main startup script (start.ts) -- an older, outdated copy",
    ]
    for place in places:
        doc.add_paragraph(place, style="List Bullet")
    doc.add_paragraph(
        "When two workers have the same name but do the job differently, chaos ensues. The program would crash "
        "because it could not decide which version to use, and the two versions expected different inputs."
    )
    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        "We simply deleted the old duplicate copy from start.ts, leaving only the official imported version. "
        "One worker, one job, no confusion."
    )

    # ---- CHALLENGE 5 ----
    doc.add_heading("11. Challenge #5: The App Was Too Slow to Wake Up (Startup Timeout)", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "Azure gives every application a limited amount of time to start up and respond to health checks. "
        "The default is about 230 seconds (~4 minutes). Our application needs to boot 9 separate programs "
        "in sequence, and on a basic (B1) tier computer, this can take longer than 4 minutes. "
        "Azure would give up waiting and declare the app 'dead' even though it was still starting."
    )
    doc.add_heading("The Solution", 2)
    doc.add_paragraph("We increased the patience setting:")
    add_code_block(doc, "WEBSITES_CONTAINER_START_TIME_LIMIT = 600")
    doc.add_paragraph(
        "This tells Azure: 'Please wait up to 10 minutes before giving up.' "
        "This gave the application plenty of time to start all 9 workers at its own pace."
    )

    # ---- CHALLENGE 6 ----
    doc.add_heading("12. Challenge #6: Secret Passwords and Security", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph(
        "Our application needs a secret password (API Key) to talk to the AI Foundry brain. "
        "If we packed this password into the ZIP file, anyone who intercepted the file could steal it "
        "and run up charges on our AI account, or worse, access sensitive data."
    )
    doc.add_heading("The Solution", 2)
    doc.add_paragraph(
        "We NEVER put passwords in the code. Instead, we store them as 'App Settings' directly inside "
        "Azure's secure configuration panel (as described in Section 5). When the application starts, "
        "it reads these settings from the local environment (like checking a locked safe in the office) "
        "rather than carrying them in the moving boxes."
    )
    doc.add_page_break()

    # ---- FINAL SETUP ----
    doc.add_heading("13. The Final Working Setup", 1)
    doc.add_paragraph("After solving all of the above challenges, here is the successfully deployed application:")

    final = [
        ("Application URL", "https://contract-agentops-eastus.azurewebsites.net/"),
        ("Health Check URL", "https://contract-agentops-eastus.azurewebsites.net/api/v1/health"),
        ("Region", "East US 2"),
        ("Computer Size", "B1 (Basic tier, Linux)"),
        ("Runtime", "Node.js 20 LTS"),
        ("AI Model", "GPT-4o via Azure AI Foundry in East US"),
        ("Status", "All 8 expert servers ONLINE, Gateway HEALTHY"),
    ]
    t4 = doc.add_table(rows=len(final) + 1, cols=2)
    t4.style = "Light Grid Accent 1"
    set_cell(t4.rows[0].cells[0], "Item", bold=True)
    set_cell(t4.rows[0].cells[1], "Value", bold=True)
    for i, (k, v) in enumerate(final, 1):
        set_cell(t4.rows[i].cells[0], k)
        set_cell(t4.rows[i].cells[1], v)

    doc.add_paragraph("")
    doc.add_paragraph("The health check returns:")
    add_code_block(doc,
        '{\n'
        '  "status": "ok",\n'
        '  "mode": "live",\n'
        '  "servers": {\n'
        '    "contract-intake-mcp": "online",\n'
        '    "contract-extraction-mcp": "online",\n'
        '    "contract-compliance-mcp": "online",\n'
        '    "contract-workflow-mcp": "online",\n'
        '    "contract-audit-mcp": "online",\n'
        '    "contract-eval-mcp": "online",\n'
        '    "contract-drift-mcp": "online",\n'
        '    "contract-feedback-mcp": "online"\n'
        '  }\n'
        '}'
    )

    # ---- HOW TO UPDATE ----
    doc.add_heading("14. How to Update the App in the Future", 1)
    doc.add_paragraph("Now that the initial setup is done, future updates are much simpler. You do NOT need to recreate the Resource Group, App Service Plan, or Web App. You only need to:")
    future_steps = [
        "Pull the latest code from GitHub: git pull origin main",
        "Re-apply known fixes if upstream reverted them (build scripts, keep-alive, tsx in dependencies)",
        "Pack a new ZIP: tar -cf deploy.zip --exclude=node_modules --exclude=.git -C AgentOps -a .",
        "Upload to Azure: az webapp deploy --name contract-agentops-eastus --resource-group contract-agentops-eastus-rg --src-path deploy.zip --type zip",
        "Wait 3-5 minutes for build + startup",
        "Verify health at /api/v1/health -- all 8 servers should show 'online'",
    ]
    for s in future_steps:
        doc.add_paragraph(s, style="List Number")

    # ---- GLOSSARY ----
    doc.add_heading("15. Glossary of Terms", 1)
    glossary = [
        ("Azure", "Microsoft's cloud computing platform -- a network of data centers around the world that you can rent computer space from."),
        ("Azure CLI (az)", "A command-line tool you install on your laptop to create and manage Azure resources by typing commands."),
        ("App Service", "An Azure product that runs web applications without you managing the underlying computer hardware."),
        ("App Service Plan", "The size and pricing tier of the computer that runs your App Service (like choosing small/medium/large)."),
        ("Resource Group", "A container (like a folder) in Azure that groups related resources together for easy management."),
        ("Tenant", "An organization's unique identity in Azure. All of a company's Azure resources live under one tenant."),
        ("Subscription", "A billing account in Azure. One tenant can have multiple subscriptions with separate budgets."),
        ("Oryx", "Azure's automatic build system that reads your project files and installs all required software libraries."),
        ("npm install", "The command that downloads and installs all the code libraries your application depends on."),
        ("ZIP Deploy / OneDeploy", "A method of deploying by uploading a compressed archive of your code to Azure."),
        ("API Key", "A secret password string used to authenticate with an external service (like AI Foundry)."),
        ("Health Check", "An automated test that asks 'Is the application alive and working?' at regular intervals."),
        ("Quota", "A limit on how much of a cloud resource you are allowed to use in a particular region."),
        ("Region", "A physical location of a Microsoft data center (e.g., East US, West Europe)."),
        ("MCP Server", "Model Context Protocol server -- a specialized mini-application that handles one specific type of work."),
        ("Gateway", "The single entry point that receives all user requests and routes them to the correct internal service."),
        ("Node.js", "The programming language runtime our application uses (JavaScript/TypeScript)."),
        ("tsx", "A tool that lets us run TypeScript code directly without a separate compilation step."),
        ("Git / GitHub", "A version control system and website for storing, sharing, and collaborating on code."),
        ("PowerShell", "The command-line interface on Windows used to type and run commands."),
        ("Stop-parsing token (--%)", "A PowerShell trick that tells it to pass the rest of the command as-is without interpreting special characters."),
    ]
    t5 = doc.add_table(rows=len(glossary) + 1, cols=2)
    t5.style = "Light Grid Accent 1"
    set_cell(t5.rows[0].cells[0], "Term", bold=True)
    set_cell(t5.rows[0].cells[1], "Definition", bold=True)
    for i, (term, defn) in enumerate(glossary, 1):
        set_cell(t5.rows[i].cells[0], term)
        set_cell(t5.rows[i].cells[1], defn)

    # Save
    doc_path = "C:\\Users\\pfernandes\\AgentOps\\AgentOps\\Contract_AgentOps_Deployment_Guide_Full.docx"
    doc.save(doc_path)
    print(f"Document created successfully at {doc_path}")

if __name__ == "__main__":
    create_doc()
