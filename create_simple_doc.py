from docx import Document
from docx.shared import Pt, Inches

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('How Our AgentOps Application is Deployed', 0)
    
    # Introduction
    doc.add_heading('1. The Basic Idea', level=1)
    p = doc.add_paragraph('Think of "deploying" an application like moving a business into a new office building. ')
    p.add_run('Right now, the application code lives on a single developer\'s laptop (like working out of a garage). ')
    p.add_run('To let everyone on the internet access it reliably, we need to rent space in a large, professional building. ')
    p.add_run('In our case, that "office building" is Microsoft Azure.')
    
    # Where it lives
    doc.add_heading('2. Where Does It Live?', level=1)
    doc.add_paragraph('We use a service called "Azure Web Apps." It is essentially a computer in the cloud that Microsoft maintains for us. We do not have to worry about power, internet cables, or hardware breaking. Microsoft handles all the physical security and maintenance.')
    
    # AI Connection
    doc.add_heading('3. Connecting to the AI Brain', level=1)
    doc.add_paragraph('The AgentOps application needs to be smart. Instead of building a new brain from scratch, it connects to a specialized "AI Brain" called Microsoft AI Foundry. We give our cloud computer a safe "key" (password) so it is allowed to ask the AI brain questions and get answers back.')
    
    # The Steps
    doc.add_heading('4. The Simple Steps to Deploy', level=1)
    doc.add_paragraph('When we have new updates or features ready for the world to see, here is how we move them to the cloud:')
    
    steps = [
        ('A. Pack the Boxes (Zipping the Files)', 'We gather all the code files and compress them into a single file, like packing up your office into moving boxes.'),
        ('B. Drive the Moving Truck (Uploading)', 'We securely upload this single packaged file from our laptop directly to our rented space in Microsoft Azure.'),
        ('C. Unpack and Set Up (Azure Build)', 'Once the package arrives, Azure automatically opens the boxes, organizes the files, and installs any extra tools the application needs to run properly.'),
        ('D. Turn on the Open Sign (Starting the App)', 'Azure starts up the application. Within a few minutes, the website comes to life and is ready for users to visit.')
    ]
    
    for step_title, step_desc in steps:
        p_step = doc.add_paragraph(style='List Bullet')
        p_step.add_run(step_title + ': ').bold = True
        p_step.add_run(step_desc)
        
    # Summary
    doc.add_heading('5. Why Do We Do It This Way?', level=1)
    p_summary = doc.add_paragraph()
    p_summary.add_run('This process is highly automated. Once the initial setup is done, updating the site is as simple as packing a new box and sending it up. If anything ever goes wrong, we can easily roll back to sending the previous box instead.')

    # Save
    doc_path = 'C:\\Users\\pfernandes\\AgentOps\\AgentOps\\Simple_Deployment_Guide.docx'
    doc.save(doc_path)
    print(f"Document created successfully at {doc_path}")

if __name__ == '__main__':
    create_doc()
